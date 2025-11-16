"""
Модуль для автоматического исправления ошибок форматирования в документах DOCX.
"""

import os
import logging
import tempfile
from pathlib import Path
from docx import Document

logger = logging.getLogger(__name__)

class CorrectionService:
    """
    Класс для автоматического исправления ошибок форматирования в документах DOCX.
    """
    
    def __init__(self):
        """Инициализация сервиса исправления документов"""
        # Настройки требований к форматированию
        self.requirements = {
            "font": "Times New Roman",
            "font_size": 14,
            "line_spacing": 1.5,
            "margins": {
                "left": 3.0,
                "right": 1.5,
                "top": 2.0,
                "bottom": 2.0
            },
            "paragraph_indent": 1.25,
            "header_formatting": {
                "bold": True,
                "alignment": "center"
            }
        }
    
    def fix_errors(self, document, errors_to_fix):
        """
        Исправляет ошибки форматирования в документе.
        
        Args:
            document (dict): Данные документа для исправления
            errors_to_fix (list): Список ошибок для исправления
            
        Returns:
            dict: Результат исправления ошибок
        """
        fixed_errors = []
        
        for error in errors_to_fix:
            error_type = error.get("type", "")
            
            # Исправление шрифта
            if error_type == "font":
                self._fix_font(document)
                fixed_errors.append(error)
            
            # Исправление размера шрифта
            elif error_type == "font_size":
                self._fix_font_size(document)
                fixed_errors.append(error)
            
            # Исправление полей страницы
            elif error_type == "margins":
                self._fix_margins(document)
                fixed_errors.append(error)
            
            # Исправление межстрочного интервала
            elif error_type == "line_spacing":
                self._fix_line_spacing(document)
                fixed_errors.append(error)
            
            # Исправление форматирования заголовков
            elif error_type == "header_format":
                self._fix_header_formatting(document)
                fixed_errors.append(error)
        
        return {
            "status": "success",
            "fixed_errors_count": len(fixed_errors),
            "fixed_errors": fixed_errors
        }
    
    def correct_document(self, file_path, errors_to_fix):
        """
        Исправляет документ и создает его исправленную версию.
        
        Args:
            file_path (str): Путь к документу для исправления
            errors_to_fix (list): Список ошибок для исправления
            
        Returns:
            str: Путь к исправленному документу
        """
        try:
            # Загрузка документа
            doc = Document(file_path)
            
            # Исправление ошибок в документе
            fixed_errors_count = 0
            
            # Исправление шрифта и размера шрифта
            if any(error.get("type") in ["font", "font_size"] for error in errors_to_fix):
                self._fix_document_font(doc)
                fixed_errors_count += 1
            
            # Исправление полей страницы
            if any(error.get("type") == "margins" for error in errors_to_fix):
                self._fix_document_margins(doc)
                fixed_errors_count += 1
            
            # Исправление межстрочного интервала
            if any(error.get("type") == "line_spacing" for error in errors_to_fix):
                self._fix_document_line_spacing(doc)
                fixed_errors_count += 1
            
            # Исправление форматирования заголовков
            if any(error.get("type") == "header_format" for error in errors_to_fix):
                self._fix_document_headers(doc)
                fixed_errors_count += 1
            
            # Сохранение исправленного документа
            base_name = os.path.basename(file_path)
            name, ext = os.path.splitext(base_name)
            corrected_file_name = f"{name}_corrected{ext}"
            corrected_file_path = os.path.join(tempfile.gettempdir(), corrected_file_name)
            
            doc.save(corrected_file_path)
            
            logger.info(f"Документ исправлен и сохранен: {corrected_file_path}")
            return corrected_file_path
        
        except Exception as e:
            logger.error(f"Ошибка при исправлении документа: {str(e)}")
            return None
    
    def _fix_font(self, document):
        """
        Исправляет шрифт в документе.
        
        Args:
            document (dict): Данные документа для исправления
        """
        required_font = self.requirements["font"]
        
        # Исправление шрифта в параграфах
        if "paragraphs" in document:
            for paragraph in document.get("paragraphs", []):
                paragraph["font"] = required_font
    
    def _fix_font_size(self, document):
        """
        Исправляет размер шрифта в документе.
        
        Args:
            document (dict): Данные документа для исправления
        """
        required_size = self.requirements["font_size"]
        
        # Исправление размера шрифта в параграфах
        if "paragraphs" in document:
            for paragraph in document.get("paragraphs", []):
                paragraph["size"] = required_size
    
    def _fix_margins(self, document):
        """
        Исправляет поля страницы в документе.
        
        Args:
            document (dict): Данные документа для исправления
        """
        required_margins = self.requirements["margins"]
        
        # Исправление полей страницы
        if "margins" in document:
            document["margins"] = required_margins.copy()
    
    def _fix_line_spacing(self, document):
        """
        Исправляет межстрочный интервал в документе.
        
        Args:
            document (dict): Данные документа для исправления
        """
        required_spacing = self.requirements["line_spacing"]
        
        # Исправление межстрочного интервала в параграфах
        if "paragraphs" in document:
            for paragraph in document.get("paragraphs", []):
                paragraph["line_spacing"] = required_spacing
    
    def _fix_header_formatting(self, document):
        """
        Исправляет форматирование заголовков в документе.
        
        Args:
            document (dict): Данные документа для исправления
        """
        header_requirements = self.requirements["header_formatting"]
        
        # Исправление форматирования заголовков
        for paragraph in document.get("paragraphs", []):
            if paragraph.get("is_header", False):
                paragraph["bold"] = header_requirements["bold"]
                paragraph["alignment"] = header_requirements["alignment"]
    
    def _fix_document_font(self, doc):
        """
        Исправляет шрифт и размер шрифта в документе Word.
        
        Args:
            doc: Документ Word для исправления
        """
        required_font = self.requirements["font"]
        required_size = self.requirements["font_size"]
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = required_font
                run.font.size = required_size
    
    def _fix_document_margins(self, doc):
        """
        Исправляет поля страницы в документе Word.
        
        Args:
            doc: Документ Word для исправления
        """
        required_margins = self.requirements["margins"]
        
        for section in doc.sections:
            section.left_margin = required_margins["left"]
            section.right_margin = required_margins["right"]
            section.top_margin = required_margins["top"]
            section.bottom_margin = required_margins["bottom"]
    
    def _fix_document_line_spacing(self, doc):
        """
        Исправляет межстрочный интервал в документе Word.
        
        Args:
            doc: Документ Word для исправления
        """
        required_spacing = self.requirements["line_spacing"]
        
        for paragraph in doc.paragraphs:
            if hasattr(paragraph, 'paragraph_format'):
                paragraph.paragraph_format.line_spacing = required_spacing
    
    def _fix_document_headers(self, doc):
        """
        Исправляет форматирование заголовков в документе Word.
        
        Args:
            doc: Документ Word для исправления
        """
        for paragraph in doc.paragraphs:
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                for run in paragraph.runs:
                    run.font.bold = True 