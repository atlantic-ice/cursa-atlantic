import os
import re
import datetime
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import shutil
from docxtpl import DocxTemplate
from docxcompose.composer import Composer

class DocumentCorrector:
    """
    Класс для исправления ошибок в документе
    """
    def __init__(self):
        # Стандартные правила для курсовых работ
        self.standard_rules = {
            'font': {
                'name': 'Times New Roman',
                'size': 14.0,  # pt
            },
            'margins': {
                'left': 3.0,  # cm
                'right': 1.0,  # cm (исправлено с 1.5 на 1.0 согласно ГОСТ)
                'top': 2.0,  # cm
                'bottom': 2.0,  # cm
            },
            'line_spacing': 1.5,
            'first_line_indent': Cm(1.25),
            'headings': {
                'h1': {
                    'font_size': 16.0,  # pt
                    'bold': True,
                    'alignment': WD_PARAGRAPH_ALIGNMENT.CENTER,
                    'all_caps': True
                },
                'h2': {
                    'font_size': 14.0,  # pt
                    'bold': True,
                    'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT
                }
            }
        }
        self.errors = []
        self.temp_files = []
    
    def __del__(self):
        """
        Деструктор для очистки временных файлов
        """
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    print(f"Удален временный файл: {temp_file}")
                    
                # Проверяем, пуста ли директория, и если да, удаляем её
                temp_dir = os.path.dirname(temp_file)
                if os.path.exists(temp_dir) and not os.listdir(temp_dir):
                    os.rmdir(temp_dir)
                    print(f"Удалена пустая временная директория: {temp_dir}")
            except Exception as e:
                print(f"Ошибка при удалении временного файла {temp_file}: {str(e)}")
    
    def correct_document(self, file_path, errors=None, out_path=None):
        """
        Исправляет ошибки в документе
        
        Args:
            file_path: Путь к файлу для исправления
            errors: Список ошибок для исправления (если None, исправляем все возможные)
            out_path: Путь для сохранения исправленного файла (если None, генерируется автоматически)
            
        Returns:
            str: Путь к исправленному файлу
        """
        self.errors = errors
        
        # Проверяем существование файла
        if not os.path.exists(file_path):
            # Пробуем добавить расширение .docx если его нет
            if not file_path.lower().endswith('.docx'):
                corrected_path = file_path + '.docx'
                if os.path.exists(corrected_path):
                    file_path = corrected_path
                else:
                    raise FileNotFoundError(f"Файл не найден: {file_path}")
            else:
                raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        try:
            # Загружаем документ
            document = Document(file_path)
            
            # Если указан путь для сохранения
            if out_path:
                # Проверяем, что путь имеет правильное расширение
                if not out_path.lower().endswith('.docx'):
                    out_path = out_path + '.docx'
                
                # Создаем директорию, если ее нет
                out_dir = os.path.dirname(out_path)
                if out_dir and not os.path.exists(out_dir):
                    os.makedirs(out_dir, exist_ok=True)
            else:
                # Если путь не указан, создаем временный файл
                temp_dir = tempfile.mkdtemp()
                file_name = os.path.basename(file_path)
                
                # Убедимся, что имя имеет расширение .docx
                if not file_name.lower().endswith('.docx'):
                    file_name = os.path.splitext(file_name)[0] + '.docx'
                
                out_path = os.path.join(temp_dir, f"corrected_{file_name}")
                self.temp_files.append(out_path)
            
            # Если список ошибок не предоставлен, исправляем все, что можем
            if errors is None:
                # Применяем базовые стили перед точечными корректировками, чтобы документ выглядел системно
                self._apply_core_styles(document)
                self._correct_all(document)
            else:
                # Исправляем только указанные ошибки
                self._correct_specific_errors(document, errors)
            
            # Сохраняем исправленный документ
            document.save(out_path)
            return out_path
            
        except Exception as e:
            print(f"Ошибка при исправлении документа: {str(e)}")
            raise
    
    def _correct_all(self, document):
        """
        Исправляет все типичные ошибки в документе
        """
        # Сначала исправляем поля страницы и базовые настройки документа
        self._correct_margins(document)
        
        # Исправляем шрифт для всего документа
        self._correct_font(document)
        
        # Исправляем межстрочный интервал
        self._correct_line_spacing(document)

        # Продвигаем псевдозаголовки (обычный текст, похожий на заголовок) в корректные стили Heading
        self._promote_pseudo_headings_to_styles(document)
        
        # Исправляем заголовки разделов (улучшенная версия)
        self._correct_section_headings(document)
        
        # Исправляем оформление таблиц
        self._correct_tables(document)
        
        # Исправляем подписи к рисункам (точки в конце)
        self._correct_images(document)
        
        # Исправляем оформление формул
        self._correct_formulas(document)
        
        # Исправляем оформление списков
        self._correct_lists(document)
        
        # Исправляем библиографические ссылки
        self._correct_bibliography_references(document)
        
        # Исправляем список литературы по ГОСТу
        self._correct_gost_bibliography(document)
        
        # Исправляем оглавление
        self._correct_toc(document)
        
        # Исправляем список сокращений и условных обозначений
        self._correct_abbreviations_list(document)
        
        # Исправляем перекрестные ссылки
        self._correct_cross_references(document)
        
        # Исправляем оформление приложений
        self._correct_appendices(document)
        
        # Исправляем оформление акцентов в тексте
        self._correct_text_accents(document)
        
        # Исправляем подстрочные ссылки
        self._correct_footnotes(document)
        
        # Исправляем нумерацию страниц
        self._correct_page_numbers(document)
        
        # ОТКЛЮЧЕНО: Исправление титульного листа (удаляло весь контент)
        # self._correct_title_page(document)
        
        # Исправляем переносы в тексте
        self._correct_hyphenation(document)
        
        # В конце применяем форматирование абзацев и выравнивание
        # для гарантии правильного форматирования всего текста
        self._correct_first_line_indent(document)
        self._correct_paragraph_alignment(document)
        self._clean_extra_blank_lines(document)
    
    def _apply_core_styles(self, document):
        """Подстраивает ключевые стили Word под нормоконтрольный стандарт."""
        try:
            normal_style = document.styles['Normal']
            self._set_style_font_defaults(normal_style, self.standard_rules['font']['size'])
            normal_style.paragraph_format.first_line_indent = Cm(1.25)
            normal_style.paragraph_format.left_indent = Cm(0)
            normal_style.paragraph_format.right_indent = Cm(0)
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(0)
            normal_style.paragraph_format.line_spacing = self.standard_rules['line_spacing']
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            normal_style.paragraph_format.keep_together = False
            normal_style.paragraph_format.keep_with_next = False
        except KeyError:
            pass

        # Нормоконтрольные интервалы для заголовков:
        # H1: перед 36 пт (≈3 одинарных), после 24 пт (≈2 одинарных)
        # H2: перед 24 пт (≈2 одинарных), после 24 пт (≈2 одинарных)
        heading_defaults = [
            ('Heading 1', self.standard_rules['headings'].get('h1', {}), Pt(36), Pt(24)),
            ('Heading 2', self.standard_rules['headings'].get('h2', {}), Pt(24), Pt(24)),
        ]

        for style_name, rule, space_before, space_after in heading_defaults:
            try:
                heading_style = document.styles[style_name]
            except KeyError:
                continue

            font_size = rule.get('font_size', self.standard_rules['font']['size'])
            bold = rule.get('bold', True)
            all_caps = rule.get('all_caps', False)
            alignment = rule.get('alignment', WD_PARAGRAPH_ALIGNMENT.LEFT)

            self._set_style_font_defaults(heading_style, font_size, bold=bold, all_caps=all_caps)

            heading_style.paragraph_format.alignment = alignment
            heading_style.paragraph_format.first_line_indent = Cm(0)
            heading_style.paragraph_format.left_indent = Cm(0)
            heading_style.paragraph_format.right_indent = Cm(0)
            heading_style.paragraph_format.space_before = space_before
            heading_style.paragraph_format.space_after = space_after
            heading_style.paragraph_format.keep_with_next = True
            heading_style.paragraph_format.keep_together = True
            # Полуторный интервал и для заголовков (требование общего межстрочного интервала)
            heading_style.paragraph_format.line_spacing = self.standard_rules['line_spacing']
            heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    def _set_style_font_defaults(self, style, font_size, *, bold=False, all_caps=False):
        """Применяет единый шрифт Times New Roman к стилю."""
        font_name = self.standard_rules['font']['name']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.all_caps = all_caps

        rPr = style._element.get_or_add_rPr()
        rPr.rFonts_ascii = font_name
        rPr.rFonts_hAnsi = font_name
        rPr.rFonts_eastAsia = font_name
        rPr.rFonts_cs = font_name

    def _clean_extra_blank_lines(self, document):
        """Удаляет лишние подряд идущие пустые абзацы, сохраняя визуальную чистоту.
        ВАЖНО: НЕ трогает таблицы - они требуют особой осторожности!
        """
        def trim_paragraph_list(paragraphs, in_table=False):
            consecutive_blank = 0
            for paragraph in list(paragraphs):
                # КРИТИЧЕСКИ ВАЖНО: в таблицах НЕ удаляем пустые параграфы!
                # Они могут быть частью структуры ячеек
                if in_table:
                    continue
                    
                if paragraph.text.strip():
                    consecutive_blank = 0
                    continue

                consecutive_blank += 1
                # Удаляем только если больше 2 подряд пустых параграфов
                if consecutive_blank > 2:
                    try:
                        element = paragraph._element
                        parent = element.getparent()
                        if parent is not None:
                            parent.remove(element)
                    except Exception as e:
                        print(f"Предупреждение: не удалось удалить пустой параграф: {str(e)}")
                        continue

        # Обрабатываем только основные параграфы документа
        trim_paragraph_list(document.paragraphs, in_table=False)

        # Таблицы НЕ трогаем - они хрупкие!
        # Комментируем опасный код:
        # for table in document.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             trim_paragraph_list(cell.paragraphs, in_table=True)

    def _promote_pseudo_headings_to_styles(self, document):
        """Находит параграфы, оформленные как заголовки вручную, и присваивает им Heading 1/2.

        Эвристики:
        - Строки вида: "ГЛАВА 1 ...", "РАЗДЕЛ 2 ..." -> Heading 1
        - Нумерация вида: "1. ..." (верхний уровень) -> Heading 1
        - Нумерация вида: "1.1 ..." (второй уровень и глубже) -> Heading 2
        - ALL-CAPS короткие слова (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, СПИСОК ЛИТЕРАТУРЫ) -> Heading 1
        
        ВАЖНО: НЕ использует paragraph.text = ... для сохранения форматирования
        """
        try:
            h1_style = None
            h2_style = None
            try:
                h1_style = document.styles['Heading 1']
            except KeyError:
                pass
            try:
                h2_style = document.styles['Heading 2']
            except KeyError:
                pass

            if not (h1_style or h2_style):
                return

            h1_keywords = {
                'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'СПИСОК ЛИТЕРАТУРЫ', 'ПРИЛОЖЕНИЯ', 'ПРИЛОЖЕНИЕ'
            }

            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))

            for paragraph in document.paragraphs:
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip()
                if not text:
                    continue
                # Уже нормальный заголовок — пропускаем
                if paragraph.style and paragraph.style.name.startswith('Heading'):
                    continue

                try:
                    lower = text.lower()

                    # ГЛАВА N / РАЗДЕЛ N (только верхний уровень)
                    if re.match(r'^(глава|раздел)\s+\d+\.?\s*', lower):
                        if h1_style:
                            paragraph.style = h1_style
                        continue

                    # Многоуровневая нумерация: 1., 1.1, 1.1.1 и т.д.
                    # Определяем уровень по количеству чисел в начале
                    m = re.match(r'^(\d+(?:\.\d+)*)(?:[\s\.:\-]+)(.*)', text)
                    if m:
                        numbering = m.group(1)
                        rest = m.group(2).strip()
                        level = numbering.count('.') + 1

                        # Выбираем стиль по уровню (1->Heading 1, 2->Heading 2, >=3->Heading 3)
                        try:
                            if level == 1 and h1_style:
                                paragraph.style = h1_style
                            elif level == 2 and h2_style:
                                paragraph.style = h2_style
                            else:
                                # Для третьего и более уровней используем Heading 3, если он есть
                                if 'Heading 3' in document.styles:
                                    paragraph.style = document.styles['Heading 3']
                        except Exception as e:
                            print(f"ОШИБКА при установке стиля заголовка: {str(e)}")

                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: НЕ используем paragraph.text = ...
                        # Вместо этого изменяем runs
                        new_text = rest
                        if new_text.endswith('.'):
                            new_text = new_text.rstrip('.')
                        
                        if new_text and len(paragraph.runs) > 0:
                            # Находим первый run с нумерацией и заменяем его текст
                            for run in paragraph.runs:
                                if numbering in run.text:
                                    run.text = run.text.replace(numbering, '', 1).lstrip('. :\-–—')
                                    break
                            
                            # Удаляем завершающую точку из последнего run
                            if paragraph.runs:
                                last_run = paragraph.runs[-1]
                                if last_run.text.endswith('.'):
                                    last_run.text = last_run.text.rstrip('.')
                        continue

                    # Нумерация "1. ..." без вложений -> Heading 1
                    if re.match(r'^\d+\.(\s+|$)', text):
                        if h1_style:
                            paragraph.style = h1_style
                        continue

                    # ALL-CAPS короткие заголовки (часто структурные части)
                    if text == text.upper() and 2 <= len(text) <= 80 and not text.endswith('.'):
                        # Неформатированные ключевые слова — точно H1
                        if text in h1_keywords and h1_style:
                            paragraph.style = h1_style
                        continue
                
                except Exception as e:
                    print(f"ОШИБКА при обработке заголовка '{text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _promote_pseudo_headings_to_styles: {str(e)}")
            import traceback
            traceback.print_exc()

    def _correct_specific_errors(self, document, errors):
        """
        Исправляет только указанные ошибки
        """
        error_types = set(error.get('type') for error in errors if 'type' in error)
        
        # Группируем ошибки по типу
        font_errors = any(et.startswith('font_') for et in error_types)
        margin_errors = any(et.endswith('_margin') for et in error_types)
        spacing_errors = 'line_spacing' in error_types
        indent_errors = 'first_line_indent' in error_types
        heading_errors = any(et.startswith('heading_') for et in error_types)
        image_errors = any(et.startswith('image_') for et in error_types)
        table_errors = any(et.startswith('table_') for et in error_types)
        page_numbers_errors = any(et.startswith('page_numbers_') for et in error_types)
        paragraph_alignment_errors = 'paragraph_alignment' in error_types
        list_errors = any(et.startswith('list_') for et in error_types)
        title_page_errors = any(et.startswith('title_page_') for et in error_types)
        
        # Исправляем соответствующие группы ошибок
        if font_errors:
            self._correct_font(document)
        if margin_errors:
            self._correct_margins(document)
        if spacing_errors:
            self._correct_line_spacing(document)
        if indent_errors:
            self._correct_first_line_indent(document)
        if heading_errors:
            self._correct_section_headings(document)
        if image_errors:
            self._correct_images(document)
        if paragraph_alignment_errors:
            self._correct_paragraph_alignment(document)
        if table_errors:
            self._correct_tables(document)
        if page_numbers_errors:
            self._correct_page_numbers(document)
        if list_errors:
            self._correct_lists(document)
        # ОТКЛЮЧЕНО: Исправление титульного листа удаляет весь контент
        # if title_page_errors:
        #     self._correct_title_page(document)
    
    def _correct_font(self, document):
        """
        Исправляет шрифт для всего документа
        Добавлена защита и обработка ошибок
        """
        try:
            # Получаем список всех параграфов внутри таблиц для особой обработки
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            for paragraph in document.paragraphs:
                # Пропускаем пустые параграфы
                if not paragraph.text.strip():
                    continue
                
                try:
                    # Определяем, является ли параграф заголовком
                    is_heading = paragraph.style.name.startswith('Heading')
                    heading_level = None
                    
                    if is_heading:
                        try:
                            heading_level = int(paragraph.style.name.replace('Heading ', ''))
                        except ValueError:
                            heading_level = None
                    
                    # Применяем соответствующий стиль шрифта
                    for run in paragraph.runs:
                        try:
                            # Устанавливаем базовый шрифт для всех элементов
                            if run.font.name != self.standard_rules['font']['name']:
                                run.font.name = self.standard_rules['font']['name']
                            
                            if is_heading and heading_level == 1:
                                # Для заголовков 1 уровня
                                if run.font.size != Pt(self.standard_rules['headings']['h1']['font_size']):
                                    run.font.size = Pt(self.standard_rules['headings']['h1']['font_size'])
                                if run.font.bold != self.standard_rules['headings']['h1']['bold']:
                                    run.font.bold = self.standard_rules['headings']['h1']['bold']
                            elif is_heading and heading_level == 2:
                                # Для заголовков 2 уровня
                                if run.font.size != Pt(self.standard_rules['headings']['h2']['font_size']):
                                    run.font.size = Pt(self.standard_rules['headings']['h2']['font_size'])
                                if run.font.bold != self.standard_rules['headings']['h2']['bold']:
                                    run.font.bold = self.standard_rules['headings']['h2']['bold']
                            else:
                                # Для обычного текста
                                if run.font.size != Pt(self.standard_rules['font']['size']):
                                    run.font.size = Pt(self.standard_rules['font']['size'])
                        
                        except Exception as e:
                            print(f"ОШИБКА при установке шрифта для run: {str(e)}")
                            continue
                
                except Exception as e:
                    print(f"ОШИБКА при обработке параграфа '{paragraph.text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_font: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _correct_margins(self, document):
        """
        Исправляет поля страницы
        """
        # Устанавливаем правильные поля для всех секций документа
        for section in document.sections:
            # Устанавливаем значения полей в сантиметрах
            section.top_margin = Cm(self.standard_rules['margins']['top'])
            section.bottom_margin = Cm(self.standard_rules['margins']['bottom'])
            section.left_margin = Cm(self.standard_rules['margins']['left'])
            section.right_margin = Cm(self.standard_rules['margins']['right'])
            
            # Устанавливаем стандартную ориентацию страницы
            section.orientation = 0  # 0 - портретная ориентация
            
            # Устанавливаем размер страницы A4
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
    
    def _correct_line_spacing(self, document):
        """
        Исправляет межстрочный интервал
        Добавлена защита от таблиц и обработка ошибок
        """
        try:
            # Получаем список всех параграфов внутри таблиц
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            for paragraph in document.paragraphs:
                # Пропускаем параграфы внутри таблиц - у них свои правила
                if id(paragraph) in table_paragraphs:
                    continue
                
                # Пропускаем пустые параграфы
                if not paragraph.text.strip():
                    continue

                try:
                    pf = paragraph.paragraph_format
                    
                    # Устанавливаем полуторный интервал (1.5) для всех абзацев, включая заголовки
                    if pf.line_spacing != self.standard_rules['line_spacing']:
                        pf.line_spacing = self.standard_rules['line_spacing']
                        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

                    # Для обычного текста сбрасываем интервалы до/после; для заголовков их задают стили
                    if not paragraph.style.name.startswith('Heading'):
                        if pf.space_before != Pt(0):
                            pf.space_before = Pt(0)
                        if pf.space_after != Pt(0):
                            pf.space_after = Pt(0)
                
                except Exception as e:
                    print(f"ОШИБКА при установке интервала для параграфа '{paragraph.text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_line_spacing: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _correct_first_line_indent(self, document):
        """
        Исправляет отступы первой строки (абзацный отступ)
        ВАЖНО: осторожно с таблицами!
        """
        # Обрабатываем основные параграфы документа
        for paragraph in document.paragraphs:
            # Пропускаем пустые параграфы и заголовки
            if not paragraph.text.strip() or paragraph.style.name.startswith('Heading'):
                continue
                
            # Пропускаем подписи к рисункам (которые должны быть без отступа)
            if paragraph.text.strip().lower().startswith(('рисунок', 'рис.')):
                # Явно устанавливаем нулевой отступ для подписей к рисункам
                paragraph.paragraph_format.first_line_indent = Cm(0)
                continue
                
            # Пропускаем заголовки таблиц (которые должны быть без отступа)
            if paragraph.text.strip().lower().startswith('таблица'):
                # Явно устанавливаем нулевой отступ для заголовков таблиц
                paragraph.paragraph_format.first_line_indent = Cm(0)
                continue
                
            # Особая обработка для элементов списков
            if re.match(r'^[•\-–—]\s', paragraph.text) or re.match(r'^\d+[.)]\s', paragraph.text):
                # Для элементов списка устанавливаем отрицательный отступ первой строки
                paragraph.paragraph_format.first_line_indent = Cm(-0.5)
                paragraph.paragraph_format.left_indent = Cm(1.0)
                continue
                
            # Принудительно устанавливаем отступ первой строки 1.25 см для остальных параграфов
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            
            # Сбрасываем другие отступы, которые могут мешать
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.paragraph_format.right_indent = Cm(0)
        
        # ОСТОРОЖНО с таблицами - минимальные изменения!
        try:
            for table in document.tables:
                for row_idx, row in enumerate(table.rows):
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Пропускаем пустые параграфы в таблицах!
                            if not paragraph.text.strip():
                                continue
                            
                            try:
                                # Для заголовков таблиц (первая строка) отступ не нужен
                                if row_idx == 0:
                                    # Только если отступ не задан явно
                                    if paragraph.paragraph_format.first_line_indent != Cm(0):
                                        paragraph.paragraph_format.first_line_indent = Cm(0)
                                else:
                                    # Для остальных строк - только если отступа нет
                                    if paragraph.paragraph_format.first_line_indent is None or paragraph.paragraph_format.first_line_indent == Cm(0):
                                        paragraph.paragraph_format.first_line_indent = Cm(1.25)
                                
                                # НЕ трогаем left/right indent в таблицах!
                                # paragraph.paragraph_format.left_indent = Cm(0)
                                # paragraph.paragraph_format.right_indent = Cm(0)
                            except Exception as e:
                                print(f"Предупреждение: ошибка установки отступа в ячейке таблицы: {str(e)}")
                                continue
        except Exception as e:
            print(f"ОШИБКА при обработке отступов в таблицах: {str(e)}")
    
    def _correct_section_headings(self, document):
        """
        Улучшенная функция для форматирования заголовков разделов
        ВАЖНО: НЕ использует paragraph.text = ... для сохранения форматирования
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            # Паттерны для идентификации заголовков разделов
            chapter_patterns = [
                r'^глава\s+\d+\.?\s+',
                r'^раздел\s+\d+\.?\s+',
                r'^\d+\.\s+[А-Я]',
                r'^\d+\.\d+\.\s+[А-Я]'
            ]
            
            # Словарь для хранения информации об уровнях заголовков
            heading_levels = {}
            
            # Первый проход: определяем уровни заголовков по нумерации
            for i, paragraph in enumerate(document.paragraphs):
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip()
                
                # Пропускаем пустые параграфы
                if not text:
                    continue
                    
                try:
                    # Проверяем, является ли параграф потенциальным заголовком
                    is_heading = False
                    heading_level = None
                    
                    # Проверяем по регулярным выражениям
                    for pattern in chapter_patterns:
                        if re.match(pattern, text, re.IGNORECASE):
                            is_heading = True
                            
                            # Определяем уровень заголовка по количеству чисел в нумерации
                            numbers = re.findall(r'\d+', text.split()[0])
                            heading_level = len(numbers)
                            break
                    
                    # Также проверяем, является ли параграф уже заголовком по стилю
                    if paragraph.style and paragraph.style.name.startswith('Heading'):
                        is_heading = True
                        try:
                            current_level = int(paragraph.style.name.replace('Heading ', ''))
                            heading_level = current_level if heading_level is None else heading_level
                        except (ValueError, AttributeError):
                            pass
                    
                    # Если это заголовок, сохраняем информацию
                    if is_heading and heading_level:
                        heading_levels[i] = heading_level
                
                except Exception as e:
                    print(f"ОШИБКА при анализе заголовка '{text[:50]}...': {str(e)}")
                    continue
            
            # Второй проход: форматируем заголовки согласно их уровню
            for i, paragraph in enumerate(document.paragraphs):
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                if i in heading_levels:
                    try:
                        level = heading_levels[i]
                        
                        # Применяем соответствующий стиль заголовка
                        if level == 1:
                            paragraph.style = document.styles['Heading 1']
                            # Дополнительное форматирование для Heading 1
                            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            paragraph.paragraph_format.space_before = Pt(12)
                            paragraph.paragraph_format.space_after = Pt(12)
                            
                            # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: делаем текст заглавными буквами через runs
                            for run in paragraph.runs:
                                run.text = run.text.upper()
                        
                        elif level == 2:
                            paragraph.style = document.styles['Heading 2']
                            # Дополнительное форматирование для Heading 2
                            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            paragraph.paragraph_format.space_before = Pt(12)
                            paragraph.paragraph_format.space_after = Pt(6)
                            
                            # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: capitalize через runs
                            # (сохраняем первое слово как есть, т.к. это может быть номер)
                            if paragraph.runs:
                                full_text = paragraph.text
                                parts = full_text.split(' ', 1)
                                if len(parts) > 1:
                                    new_text = parts[0] + ' ' + parts[1].capitalize()
                                    # Применяем к первому run
                                    for run in paragraph.runs:
                                        run.text = ''
                                    if paragraph.runs:
                                        paragraph.runs[0].text = new_text
                        
                        elif level == 3:
                            paragraph.style = document.styles['Heading 3']
                            # Дополнительное форматирование для Heading 3
                            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            paragraph.paragraph_format.space_before = Pt(6)
                            paragraph.paragraph_format.space_after = Pt(3)
                        
                        # Общие правила для всех заголовков
                        paragraph.paragraph_format.first_line_indent = Cm(0)
                        paragraph.paragraph_format.left_indent = Cm(0)
                        paragraph.paragraph_format.right_indent = Cm(0)
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: убираем точку через runs
                        if paragraph.text.strip().endswith('.'):
                            if paragraph.runs:
                                last_run = paragraph.runs[-1]
                                last_run.text = last_run.text.rstrip('.')
                        
                        # Форматирование шрифта заголовка
                        for run in paragraph.runs:
                            run.font.name = self.standard_rules['font']['name']
                            if level == 1:
                                run.font.size = Pt(16)
                                run.font.bold = True
                            elif level == 2:
                                run.font.size = Pt(14)
                                run.font.bold = True
                            else:
                                run.font.size = Pt(14)
                                run.font.bold = True
                    
                    except Exception as e:
                        print(f"ОШИБКА при форматировании заголовка '{paragraph.text[:50]}...': {str(e)}")
                        continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_section_headings: {str(e)}")
            import traceback
            traceback.print_exc() 

    def _correct_images(self, document):
        """
        Исправляет подписи к рисункам
        ВАЖНО: НЕ использует paragraph.text = ... для сохранения форматирования
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            for paragraph in document.paragraphs:
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip()
                if not text:
                    continue
                
                try:
                    # Проверяем, является ли параграф подписью к рисунку
                    if text.lower().startswith(('рис.', 'рисунок', 'рис ')):
                        # Заменяем сокращение "рис." на полное "Рисунок"
                        if text.lower().startswith('рис.') or text.lower().startswith('рис '):
                            number_match = re.search(r'рис\.?\s*(\d+)', text.lower())
                            if number_match:
                                number = number_match.group(1)
                                text_after = text[number_match.end():].lstrip()
                                
                                # Проверяем, есть ли разделитель между номером и названием
                                if text_after.startswith('-'):
                                    text_after = text_after[1:].lstrip()
                                elif not text_after.startswith('–') and not text_after.startswith('—'):
                                    text_after = '– ' + text_after
                                
                                # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: изменяем через runs, а не paragraph.text
                                new_text = f"Рисунок {number} {text_after}"
                                
                                # Очищаем все runs и создаём один новый с правильным текстом
                                # Сохраняем форматирование первого run
                                if paragraph.runs:
                                    first_run = paragraph.runs[0]
                                    # Сохраняем форматирование
                                    font_name = first_run.font.name
                                    font_size = first_run.font.size
                                    
                                    # Очищаем все runs
                                    for run in paragraph.runs:
                                        run.text = ''
                                    
                                    # Устанавливаем новый текст в первый run
                                    paragraph.runs[0].text = new_text
                                    paragraph.runs[0].font.name = font_name
                                    if font_size:
                                        paragraph.runs[0].font.size = font_size
                        
                        # Добавляем точку в конце подписи, если её нет
                        if not paragraph.text.strip().endswith('.'):
                            if paragraph.runs:
                                last_run = paragraph.runs[-1]
                                last_run.text = last_run.text.rstrip() + '.'
                        
                        # Выравниваем подпись по центру
                        pf = paragraph.paragraph_format
                        if pf.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Сбрасываем отступы
                        if pf.first_line_indent != Cm(0):
                            pf.first_line_indent = Cm(0)
                        if pf.left_indent != Cm(0):
                            pf.left_indent = Cm(0)
                
                except Exception as e:
                    print(f"ОШИБКА при обработке подписи к рисунку '{text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_images: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _correct_paragraph_alignment(self, document):
        """
        Исправляет выравнивание параграфов
        ВАЖНО: НЕ трогаем параграфы в таблицах - они обрабатываются отдельно!
        """
        # Сначала проходим по всем параграфам в основном документе
        for paragraph in document.paragraphs:
            # Пропускаем пустые параграфы
            if not paragraph.text.strip():
                continue
                
            # Обрабатываем заголовки
            if paragraph.style.name.startswith('Heading'):
                heading_level = int(paragraph.style.name.replace('Heading ', ''))
                if heading_level == 1:
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                continue
                
            # Подписи к рисункам выравниваем по центру
            if paragraph.text.strip().lower().startswith(('рисунок', 'рис.')):
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                continue
                
            # Заголовки таблиц выравниваем по левому краю
            if paragraph.text.strip().lower().startswith('таблица'):
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                continue
                
            # Элементы списков выравниваем по ширине, но с особым форматированием
            if re.match(r'^[•\-–—]\s', paragraph.text) or re.match(r'^\d+[.)]\s', paragraph.text):
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                continue
                
            # Все остальные параграфы (основной текст) выравниваем по ширине
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Включаем автоматические переносы для улучшения выравнивания по ширине
            self._enable_hyphenation(paragraph)
        
        # Затем проходим по всем таблицам и выравниваем текст в ячейках
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Пропускаем пустые параграфы
                        if not paragraph.text.strip():
                            continue
                        
                        # Выравниваем текст в ячейках по ширине
                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY                        # Включаем автоматические переносы для улучшения выравнивания
                        self._enable_hyphenation(paragraph)

    def _enable_hyphenation(self, paragraph):
        """
        Включает автоматические переносы для параграфа
        """
        try:
            pPr = paragraph._element.get_or_add_pPr()
            if pPr is not None:
                # Добавляем свойство автоматического переноса слов
                hyphenation_element = OxmlElement('w:autoSpaceDE')
                hyphenation_element.set(qn('w:val'), '1')
                pPr.append(hyphenation_element)
                
                # Добавляем свойство выравнивания последней строки
                last_line_element = OxmlElement('w:contextualSpacing')
                last_line_element.set(qn('w:val'), '1')
                pPr.append(last_line_element)
        except Exception as e:
            print(f"Предупреждение: Не удалось включить переносы слов: {str(e)}")
    
    def _correct_tables(self, document):
        """
        Исправляет оформление таблиц с защитой структуры
        """
        try:
            for table in document.tables:
                # Сохраняем исходные свойства таблицы
                try:
                    table_alignment = table.alignment
                except:
                    table_alignment = WD_TABLE_ALIGNMENT.LEFT
                
                # Обрабатываем каждую ячейку осторожно
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        # Сохраняем объединения ячеек - НЕ трогаем merged cells
                        try:
                            # Проверяем, не является ли ячейка частью объединения
                            tc = cell._element
                            tcPr = tc.get_or_add_tcPr()
                            vMerge = tcPr.find(qn('w:vMerge'))
                            hMerge = tcPr.find(qn('w:hMerge'))
                            gridSpan = tcPr.find(qn('w:gridSpan'))
                            
                            # Если ячейка объединена - минимальные изменения
                            is_merged = (vMerge is not None or hMerge is not None or gridSpan is not None)
                        except:
                            is_merged = False
                        
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            # КРИТИЧЕСКИ ВАЖНО: не удаляем пустые параграфы в таблицах!
                            # Они могут быть частью структуры
                            
                            # Применяем форматирование только к непустым параграфам
                            if paragraph.text.strip():
                                try:
                                    # Устанавливаем шрифт
                                    for run in paragraph.runs:
                                        if run.font.name != self.standard_rules['font']['name']:
                                            run.font.name = self.standard_rules['font']['name']
                                        if run.font.size != Pt(self.standard_rules['font']['size']):
                                            run.font.size = Pt(self.standard_rules['font']['size'])
                                    
                                    # Выравнивание: только если не задано явно
                                    if paragraph.paragraph_format.alignment is None or paragraph.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
                                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                    
                                    # Отступ: НЕ трогаем в заголовках (первая строка) и объединённых ячейках
                                    if row_idx > 0 and not is_merged:
                                        # Проверяем, не задан ли уже отступ
                                        if paragraph.paragraph_format.first_line_indent is None or paragraph.paragraph_format.first_line_indent == Cm(0):
                                            paragraph.paragraph_format.first_line_indent = Cm(1.25)
                                    elif row_idx == 0:
                                        # Для первой строки (заголовок) - явно убираем отступ
                                        paragraph.paragraph_format.first_line_indent = Cm(0)
                                    
                                    # Межстрочный интервал - аккуратно
                                    if paragraph.paragraph_format.line_spacing_rule != WD_LINE_SPACING.MULTIPLE:
                                        paragraph.paragraph_format.line_spacing = self.standard_rules['line_spacing']
                                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                                    
                                    # Убираем интервалы ДО и ПОСЛЕ только внутри ячеек
                                    paragraph.paragraph_format.space_before = Pt(0)
                                    paragraph.paragraph_format.space_after = Pt(0)
                                    
                                except Exception as e:
                                    # Логируем ошибку, но продолжаем обработку
                                    print(f"Предупреждение: ошибка форматирования параграфа в ячейке [{row_idx}][{cell_idx}]: {str(e)}")
                                    continue
                
        except Exception as e:
            print(f"ОШИБКА при исправлении таблиц: {str(e)}")
            # Не падаем, просто логируем и продолжаем
        
        # Исправляем заголовки таблиц (вне таблицы)
        try:
            for paragraph in document.paragraphs:
                text_lower = paragraph.text.strip().lower()
                if text_lower.startswith('таблица'):
                    # Форматирование заголовка таблицы
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    paragraph.paragraph_format.space_after = Pt(6)
                    paragraph.paragraph_format.space_before = Pt(12)
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    
                    # Добавляем точку в конце подписи, если её нет
                    if not paragraph.text.strip().endswith('.'):
                        paragraph.text = paragraph.text.strip() + '.'
        except Exception as e:
            print(f"ОШИБКА при исправлении заголовков таблиц: {str(e)}")
    
    def _correct_page_numbers(self, document):
        """
        Исправляет нумерацию страниц (добавляет нумерацию в верхний колонтитул справа)
        """
        # Для каждой секции документа
        for section in document.sections:
            # Отключаем связь с предыдущим колонтитулом
            section.header.is_linked_to_previous = False
            
            # Получаем доступ к верхнему колонтитулу
            header = section.header
            
            # Очищаем содержимое верхнего колонтитула
            for paragraph in header.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
            
            # Создаем новый параграф для номера страницы
            header_paragraph = header.add_paragraph()
            
            # Настраиваем форматирование параграфа
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Выравнивание справа
            header_paragraph.paragraph_format.space_before = Pt(0)
            header_paragraph.paragraph_format.space_after = Pt(0)
            header_paragraph.paragraph_format.line_spacing = 1.0
            
            # Добавляем номер страницы через поле
            run = header_paragraph.add_run()
            self._add_page_number(run)
            
            # Устанавливаем шрифт для номера страницы
            run.font.name = self.standard_rules['font']['name']  # Times New Roman
            run.font.size = Pt(12)  # 12pt
            run.font.bold = False
            
            # Добавляем пустой run после номера страницы для корректного отображения
            header_paragraph.add_run()
        
        # Отключаем нумерацию на первых страницах
        self._suppress_initial_page_numbers(document)
    
    def _add_page_number(self, run):
        """
        Добавляет номер страницы через поле
        """
        # Создаем элементы для номера страницы
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE \\* MERGEFORMAT "  # Добавляем MERGEFORMAT для лучшей совместимости

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "1"  # Placeholder для номера страницы

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        # Добавляем созданные элементы в параграф
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(fldChar3)
        run._element.append(fldChar4)
    
    def _suppress_initial_page_numbers(self, document):
        """
        Отключает нумерацию на начальных страницах (титульный лист, задание, реферат, оглавление)
        и устанавливает начало нумерации с нужной страницы
        """
        try:
            if document.sections:
                first_section = document.sections[0]
                
                # Получаем XML-элемент секции
                section_props = first_section._sectPr
                
                # Создаем новый элемент titlePg
                title_pg = OxmlElement('w:titlePg')
                
                # Проверяем, существует ли уже элемент
                if section_props.find(qn('w:titlePg')) is None:
                    section_props.append(title_pg)
                
                # Устанавливаем начальный номер страницы (3 или 4, в зависимости от оглавления)
                # Находим элемент pgNumType или создаем новый
                pg_num_type = section_props.find(qn('w:pgNumType'))
                if pg_num_type is None:
                    pg_num_type = OxmlElement('w:pgNumType')
                    section_props.append(pg_num_type)
                
                # Устанавливаем начальный номер страницы как 3
                # (можно будет вручную изменить на 4, если оглавление занимает две страницы)
                pg_num_type.set(qn('w:start'), '3')
                
        except Exception as e:
            print(f"Предупреждение: Не удалось настроить нумерацию начальных страниц: {str(e)}")
            # Продолжаем выполнение даже при ошибке, т.к. это некритичная функция
    
    def _correct_lists(self, document):
        """
        Исправляет оформление списков
        ВАЖНО: обрабатывает только списки вне таблиц
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            for paragraph in document.paragraphs:
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                # Пропускаем пустые параграфы и заголовки
                if not paragraph.text.strip() or paragraph.style.name.startswith('Heading'):
                    continue
                
                try:
                    # Проверяем, является ли параграф элементом списка
                    is_list_item = False
                    list_type = None
                    
                    # Проверка на маркированный список
                    if re.match(r'^[•\-–—]\s', paragraph.text):
                        is_list_item = True
                        list_type = 'bullet'
                    
                    # Проверка на нумерованный список
                    elif re.match(r'^\d+[.)]\s', paragraph.text) or re.match(r'^[a-яa-z][.)]\s', paragraph.text, re.IGNORECASE):
                        is_list_item = True
                        list_type = 'numbered'
                    
                    # Применяем форматирование к элементам списка
                    if is_list_item:
                        pf = paragraph.paragraph_format
                        
                        # БЕЗОПАСНАЯ УСТАНОВКА: проверяем текущее значение перед изменением
                        if pf.left_indent is None or pf.left_indent != Cm(1.0):
                            pf.left_indent = Cm(1.0)
                        
                        if pf.first_line_indent is None or pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)  # Обратный отступ для маркера
                        
                        # Устанавливаем межстрочный интервал
                        if pf.line_spacing != self.standard_rules['line_spacing']:
                            pf.line_spacing = self.standard_rules['line_spacing']
                            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        
                        # Выравнивание по ширине только если не установлено
                        if pf.alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        
                        # Шрифт для элементов списка - БЕЗ УДАЛЕНИЯ RUNS
                        for run in paragraph.runs:
                            if run.font.name != self.standard_rules['font']['name']:
                                run.font.name = self.standard_rules['font']['name']
                            if run.font.size != Pt(self.standard_rules['font']['size']):
                                run.font.size = Pt(self.standard_rules['font']['size'])
                
                except Exception as e:
                    print(f"ОШИБКА при обработке элемента списка '{paragraph.text[:50]}...': {str(e)}")
                    continue
            
            # Дополнительно обрабатываем буквенные перечисления
            self._correct_letter_lists(document)
            
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_lists: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _correct_letter_lists(self, document):
        """
        Исправляет оформление перечислений с буквенной нумерацией
        ВАЖНО: сохраняет форматирование текста, не перезаписывает paragraph.text
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            # Регулярное выражение для поиска буквенных перечислений
            letter_list_pattern = r'^([а-яa-z])[)\.]\s'
            
            for paragraph in document.paragraphs:
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip()
                
                # Пропускаем пустые параграфы и заголовки
                if not text or paragraph.style.name.startswith('Heading'):
                    continue
                
                try:
                    # Проверяем, является ли параграф элементом буквенного перечисления
                    match = re.match(letter_list_pattern, text, re.IGNORECASE)
                    if match:
                        pf = paragraph.paragraph_format
                        
                        # Форматируем перечисление - БЕЗОПАСНО
                        if pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)  # Обратный отступ для буквы
                        
                        if pf.left_indent != Cm(1.0):
                            pf.left_indent = Cm(1.0)  # Отступ слева для текста
                        
                        if pf.alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        
                        # Проверяем формат буквы - должен быть с закрывающей скобкой: а)
                        letter = match.group(1)
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: НЕ ИСПОЛЬЗУЕМ paragraph.text = ...
                        # Вместо этого изменяем только первый run
                        if text[1] == '.' and len(paragraph.runs) > 0:
                            # Ищем точку в первом run и заменяем на скобку
                            first_run = paragraph.runs[0]
                            if '.' in first_run.text:
                                first_run.text = first_run.text.replace(f"{letter}.", f"{letter})", 1)
                        
                        # Восстанавливаем форматирование шрифта БЕЗ УНИЧТОЖЕНИЯ runs
                        for run in paragraph.runs:
                            if run.font.name != self.standard_rules['font']['name']:
                                run.font.name = self.standard_rules['font']['name']
                            if run.font.size != Pt(self.standard_rules['font']['size']):
                                run.font.size = Pt(self.standard_rules['font']['size'])
                
                except Exception as e:
                    print(f"ОШИБКА при обработке буквенного перечисления '{text[:50]}...': {str(e)}")
                    continue
            
            # УЛУЧШЕННАЯ ЛОГИКА многоуровневых перечислений
            self._correct_multilevel_lists(document, table_paragraphs)
            
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_letter_lists: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _correct_multilevel_lists(self, document, table_paragraphs):
        """
        Форматирует многоуровневые перечисления с правильными отступами
        """
        try:
            # Регулярные выражения для разных уровней
            level1_pattern = r'^(\d+)[)\.]\s'  # 1) или 1.
            level2_pattern = r'^([а-яa-z])[)\.]\s'  # а) или а.
            level3_pattern = r'^[•\-–—]\s'  # маркеры
            
            in_list = False
            
            for paragraph in document.paragraphs:
                # Пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip()
                
                # Пропускаем пустые параграфы и заголовки
                if not text or paragraph.style.name.startswith('Heading'):
                    in_list = False
                    continue
                
                try:
                    pf = paragraph.paragraph_format
                    
                    # Проверяем уровень списка
                    if re.match(level1_pattern, text):
                        # Начало нумерованного списка (первый уровень)
                        in_list = True
                        
                        # Форматируем первый уровень - БЕЗОПАСНО
                        if pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)
                        if pf.left_indent != Cm(0.5):
                            pf.left_indent = Cm(0.5)
                        
                    elif re.match(level2_pattern, text, re.IGNORECASE) and in_list:
                        # Буквенное перечисление (второй уровень)
                        
                        # Форматируем второй уровень с дополнительным отступом - БЕЗОПАСНО
                        if pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)
                        if pf.left_indent != Cm(1.5):
                            pf.left_indent = Cm(1.5)  # Увеличенный отступ для вложенного списка
                        
                    elif re.match(level3_pattern, text) and in_list:
                        # Маркированный список (третий уровень)
                        
                        # Форматируем третий уровень - БЕЗОПАСНО
                        if pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)
                        if pf.left_indent != Cm(2.5):
                            pf.left_indent = Cm(2.5)  # Еще больший отступ
                        
                    elif in_list and text:
                        # Обычный параграф между элементами списка - НЕ СБРАСЫВАЕМ in_list
                        # Проверяем, является ли это продолжением элемента списка
                        if not (re.match(level1_pattern, text) or 
                                re.match(level2_pattern, text, re.IGNORECASE) or 
                                re.match(level3_pattern, text)):
                            # Это обычный текст - конец списка только если нет отступа
                            if pf.left_indent is None or pf.left_indent < Cm(0.5):
                                in_list = False
                
                except Exception as e:
                    print(f"ОШИБКА при форматировании многоуровневого списка '{text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_multilevel_lists: {str(e)}")
            import traceback
            traceback.print_exc() 

    def _correct_title_page(self, document):
        """
        Исправляет титульный лист: заменяет неправильный титульный лист на шаблонный
        или вставляет шаблонный титульный лист, если он отсутствует
        """
        # Проверяем наличие титульного листа
        title_page_exists = False
        title_page_paragraphs = []
        
        # Собираем все параграфы титульного листа (до первого Heading 1 или до "СОДЕРЖАНИЕ"/"ВВЕДЕНИЕ")
        for i, para in enumerate(document.paragraphs):
            text = para.text.strip().lower()
            if para.style.name.startswith('Heading') and para.style.name == 'Heading 1':
                break
            if any(word in text for word in ['содержание', 'введение']):
                break
            title_page_paragraphs.append((i, para))
        
        # Определяем, содержит ли документ титульный лист
        title_keywords = [
            'университет', 'кафедра', 'факультет', 'курсовая', 'работа', 'студент', 'руководитель'
        ]
        
        # Проверяем, содержит ли титульный лист ключевые слова
        if len(title_page_paragraphs) > 3:  # Минимальное количество строк для титульного листа
            title_text = ' '.join([para[1].text.lower() for para in title_page_paragraphs])
            if any(keyword in title_text for keyword in title_keywords):
                title_page_exists = True
        
        # Путь к шаблону титульного листа
        template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'templates', 'Титульный лист.docx')
        
        if not os.path.exists(template_path):
            print(f"Предупреждение: Шаблон титульного листа не найден по пути {template_path}")
            return

        if title_page_exists:
            # Если титульный лист существует, но неправильно оформлен - заменяем его
            # Удаляем старые параграфы титульного листа
            for i, para in reversed(title_page_paragraphs):
                p = para._element
                p.getparent().remove(p)
                
            # Вставляем новый титульный лист из шаблона
            self._insert_title_page_from_template(document, template_path)
        else:
            # Если титульного листа нет - просто вставляем шаблон
            self._insert_title_page_from_template(document, template_path)

    def _insert_title_page_from_template(self, document, template_path):
        """
        Вставляет титульный лист из шаблона в начало документа
        """
        try:
            # Создаем временный файл для сохранения текущего документа
            temp_doc_fd, temp_doc_path = tempfile.mkstemp(suffix='.docx')
            os.close(temp_doc_fd)
            document.save(temp_doc_path)
            self.temp_files.append(temp_doc_path)
            
            # Создаем новый документ на основе шаблона титульного листа
            title_doc = Document(template_path)
            
            # Сохраняем шаблон во временный файл
            temp_title_fd, temp_title_path = tempfile.mkstemp(suffix='.docx')
            os.close(temp_title_fd)
            title_doc.save(temp_title_path)
            self.temp_files.append(temp_title_path)
            
            # Используем Composer для объединения документов
            composer = Composer(title_doc)
            doc_to_merge = Document(temp_doc_path)
            composer.append(doc_to_merge)
            
            # Сохраняем результат во временный файл
            temp_result_fd, temp_result_path = tempfile.mkstemp(suffix='.docx')
            os.close(temp_result_fd)
            composer.save(temp_result_path)
            self.temp_files.append(temp_result_path)
            
            # Загружаем результат обратно в исходный документ
            result_doc = Document(temp_result_path)
            
            # Очищаем исходный документ
            for element in list(document._element.body):
                document._element.body.remove(element)
            
            # Копируем все элементы из результата в исходный документ
            for element in list(result_doc._element.body):
                document._element.body.append(element)
                
            print("Титульный лист успешно вставлен")
            
        except Exception as e:
            print(f"Ошибка при вставке титульного листа: {str(e)}")
            # Продолжаем выполнение даже при ошибке, т.к. титульный лист не критичен 

    def _correct_formulas(self, document):
        """
        Исправляет оформление формул в документе
        ОТКЛЮЧЕНО создание таблиц - слишком опасно для структуры документа
        Теперь только форматирует существующие формулы
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            for paragraph in document.paragraphs:
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip()
                if not text:
                    continue
                
                try:
                    # Ищем параграфы, которые могут содержать формулы
                    if re.search(r'(?:\(\d+\)|\(\d+\.\d+\))', text):  # Формат (1) или (1.1)
                        pf = paragraph.paragraph_format
                        
                        # Выравниваем формулы по центру
                        if pf.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Убираем абзацный отступ
                        if pf.first_line_indent != Cm(0):
                            pf.first_line_indent = Cm(0)
                        
                        # Добавляем интервал до и после формулы
                        if pf.space_before != Pt(6):
                            pf.space_before = Pt(6)
                        if pf.space_after != Pt(6):
                            pf.space_after = Pt(6)
                        
                        # ОТКЛЮЧЕНО: создание таблиц для размещения номера формулы
                        # Это было слишком опасно - удаляло параграфы и могло сломать структуру
                        # Пользователь должен сам разместить номер формулы справа через табуляцию
                
                except Exception as e:
                    print(f"ОШИБКА при форматировании формулы '{text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_formulas: {str(e)}")
            import traceback
            traceback.print_exc()

    def _correct_bibliography_references(self, document):
        """
        Исправляет оформление библиографических ссылок в тексте
        ВАЖНО: НЕ использует paragraph.text = ... для сохранения форматирования
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            # Регулярные выражения для поиска и замены
            replacements = [
                # Ссылки с круглыми скобками вместо квадратных
                (r'\((\d+)\)', r'[\1]'),
                # Ссылки с буквой "с" вне скобок
                (r'\[(\d+)\]\s*с\.\s*(\d+)', r'[\1, с. \2]'),
                # Ссылки без пробела после запятой
                (r'\[(\d+),(\d+)\]', r'[\1, \2]')
            ]
            
            # Проходим по всем параграфам документа
            for paragraph in document.paragraphs:
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                # Пропускаем пустые параграфы
                if not paragraph.text.strip():
                    continue
                
                try:
                    # Проверяем, нужны ли исправления
                    text = paragraph.text
                    needs_correction = False
                    
                    for pattern, _ in replacements:
                        if re.search(pattern, text):
                            needs_correction = True
                            break
                    
                    if needs_correction and paragraph.runs:
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: работаем через runs
                        # Применяем замены к тексту каждого run
                        for run in paragraph.runs:
                            if run.text:
                                new_text = run.text
                                for pattern, replacement in replacements:
                                    new_text = re.sub(pattern, replacement, new_text)
                                run.text = new_text
                
                except Exception as e:
                    print(f"ОШИБКА при исправлении библиографических ссылок '{paragraph.text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_bibliography_references: {str(e)}")
            import traceback
            traceback.print_exc() 

    def _correct_gost_bibliography(self, document):
        """
        Исправляет оформление списка литературы в соответствии с ГОСТ 7.0.100-2018
        ВАЖНО: НЕ использует paragraph.text = ... для сохранения форматирования
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            # Ищем раздел со списком литературы
            bibliography_started = False
            bibliography_paragraphs = []
            
            # Паттерны для идентификации различных типов источников
            source_patterns = {
                'book': r'(?i)(?:^|\s)[а-яА-Я][а-яА-Я\s]+\s[А-Я]\.\s?[А-Я]\.(?:\s|,)',
                'article': r'(?i)(?:^|\s)[а-яА-Я][а-яА-Я\s]+\s[А-Я]\.\s?[А-Я]\.\s[А-Я][а-я]',
                'web': r'(?i)(?:https?://|www\.|электронный|ресурс|доступ)',
                'law': r'(?i)(?:федеральный закон|приказ|постановление|распоряжение|указ)',
                'gost': r'(?i)(?:^|\s)ГОСТ\s',
                'dissertation': r'(?i)(?:дис|автореф)\.(?:\s|\.)'
            }
            
            for i, paragraph in enumerate(document.paragraphs):
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip().lower()
                
                try:
                    # Определяем начало списка литературы
                    if not bibliography_started and re.search(r'список\s+(использованн(ой|ых)\s+)?литератур', text):
                        bibliography_started = True
                        # Форматируем заголовок списка литературы
                        paragraph.style = document.styles['Heading 1']
                        pf = paragraph.paragraph_format
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        pf.first_line_indent = Cm(0)
                        pf.space_after = Pt(12)
                        pf.space_before = Pt(12)
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: убираем точку через runs
                        if paragraph.text.strip().endswith('.'):
                            if paragraph.runs:
                                last_run = paragraph.runs[-1]
                                last_run.text = last_run.text.rstrip('.')
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: upper через runs
                        for run in paragraph.runs:
                            run.text = run.text.upper()
                        continue
                    
                    # Собираем параграфы списка литературы
                    if bibliography_started:
                        # Если встретили новый заголовок, значит список литературы закончился
                        if paragraph.style.name.startswith('Heading'):
                            break
                        
                        # Добавляем параграф в список для дальнейшей обработки
                        bibliography_paragraphs.append((i, paragraph))
                
                except Exception as e:
                    print(f"ОШИБКА при поиске списка литературы '{text[:50]}...': {str(e)}")
                    continue
            
            # Если нашли список литературы, форматируем его
            if bibliography_paragraphs:
                for idx, (paragraph_index, paragraph) in enumerate(bibliography_paragraphs):
                    # Пропускаем пустые параграфы
                    if not paragraph.text.strip():
                        continue
                    
                    try:
                        # Определяем тип источника
                        source_type = self._detect_source_type(paragraph.text, source_patterns)
                        
                        # Форматируем элемент списка литературы
                        pf = paragraph.paragraph_format
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        pf.first_line_indent = Cm(-0.75)  # Обратный отступ для номера
                        pf.left_indent = Cm(0.75)  # Отступ слева для текста
                        pf.space_after = Pt(6)
                        pf.space_before = Pt(0)
                        
                        # Проверяем, есть ли номер в начале строки
                        numbered = re.match(r'^\d+\.?\s', paragraph.text.strip())
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: если номера нет, добавляем через runs
                        if not numbered and paragraph.runs:
                            # Добавляем номер в начало первого run
                            first_run = paragraph.runs[0]
                            first_run.text = f"{idx + 1}. {first_run.text}"
                        
                        # Применяем форматирование шрифта
                        for run in paragraph.runs:
                            if run.font.name != self.standard_rules['font']['name']:
                                run.font.name = self.standard_rules['font']['name']
                            if run.font.size != Pt(self.standard_rules['font']['size']):
                                run.font.size = Pt(self.standard_rules['font']['size'])
                    
                    except Exception as e:
                        print(f"ОШИБКА при форматировании элемента библиографии '{paragraph.text[:50]}...': {str(e)}")
                        continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_gost_bibliography: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _detect_source_type(self, text, patterns):
        """
        Определяет тип источника по его тексту
        """
        for source_type, pattern in patterns.items():
            if re.search(pattern, text):
                return source_type
        return None
    
    def _format_source_by_gost(self, text, source_type):
        """
        Форматирует библиографическую запись по ГОСТу в зависимости от типа источника
        """
        # Убираем лишние пробелы
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Проверяем, есть ли номер в начале строки и сохраняем его
        number_match = re.match(r'^(\d+\.?\s+)', text)
        number_prefix = number_match.group(1) if number_match else ""
        
        # Убираем номер для обработки и добавим его в конце
        if number_prefix:
            text = text[len(number_prefix):].strip()
        
        # Форматируем по типу источника
        if source_type == 'book':
            # Форматирование для книг
            # Фамилия, И. О. Название книги / И. О. Фамилия. – Город : Издательство, Год. – N с.
            
            # Ищем автора и название
            author_match = re.match(r'([а-яА-Я][а-яА-Я\s]+)\s([А-Я])\.\s?([А-Я])\.', text)
            if author_match:
                surname, initial1, initial2 = author_match.groups()
                rest_of_text = text[author_match.end():].strip()
                
                # Проверяем, есть ли ":" или ";" после названия для определения издательства
                publisher_match = re.search(r'[:;]\s*([^,.]+),\s*(\d{4})', rest_of_text)
                
                if publisher_match:
                    publisher, year = publisher_match.groups()
                    
                    # Форматируем строку по ГОСТу
                    formatted_text = f"{surname}, {initial1}. {initial2}. {rest_of_text}"
                    
                    # Проверяем, есть ли информация о страницах
                    if not re.search(r'\d+\s*с\.', formatted_text):
                        formatted_text += " – Текст : непосредственный."
                else:
                    # Если не удалось выделить все компоненты, оставляем как есть
                    formatted_text = text
            else:
                formatted_text = text
                
        elif source_type == 'article':
            # Форматирование для статей
            # Фамилия, И. О. Название статьи / И. О. Фамилия // Название журнала. – Год. – Т. X, № N. – С. XX-XX.
            
            # Ищем автора и название
            author_match = re.match(r'([а-яА-Я][а-яА-Я\s]+)\s([А-Я])\.\s?([А-Я])\.', text)
            if author_match:
                surname, initial1, initial2 = author_match.groups()
                rest_of_text = text[author_match.end():].strip()
                
                # Проверяем, есть ли "//" для разделения названия статьи и журнала
                journal_match = re.search(r'//\s*([^.]+)', rest_of_text)
                
                if journal_match:
                    # Форматируем строку по ГОСТу
                    formatted_text = f"{surname}, {initial1}. {initial2}. {rest_of_text}"
                else:
                    # Если не удалось выделить все компоненты, оставляем как есть
                    formatted_text = text
            else:
                formatted_text = text
                
        elif source_type == 'web':
            # Форматирование для веб-ресурсов
            # Название сайта : сайт. – URL: http://... (дата обращения: ДД.ММ.ГГГГ). – Текст : электронный.
            
            # Проверяем, есть ли URL в тексте
            url_match = re.search(r'(https?://[^\s]+)', text)
            
            if url_match:
                url = url_match.group(1)
                
                # Проверяем, есть ли фраза "дата обращения"
                if not re.search(r'дата\s+обращения', text, re.IGNORECASE):
                    # Если нет, добавляем текущую дату
                    today = datetime.datetime.now().strftime("%d.%m.%Y")
                    text = text.replace(url, f"{url} (дата обращения: {today})")
                
                # Проверяем, есть ли указание на электронный ресурс
                if not re.search(r'электронный|ресурс', text, re.IGNORECASE):
                    text += " – Текст : электронный."
                
                formatted_text = text
            else:
                formatted_text = text
                
        elif source_type == 'law':
            # Форматирование для нормативных документов
            # Об образовании в Российской Федерации : Федеральный закон № 273-ФЗ : [принят Государственной думой 21 декабря 2012 года : одобрен Советом Федерации 26 декабря 2012 года]. – Москва : Проспект, 2020. – 192 с.
            
            formatted_text = text
            
        elif source_type == 'gost':
            # Форматирование для ГОСТов
            # ГОСТ Р 7.0.100-2018. Библиографическая запись. Библиографическое описание. Общие требования и правила составления : национальный стандарт Российской Федерации : издание официальное : утвержден и введен в действие Приказом Федерального агентства по техническому регулированию и метрологии от 3 декабря 2018 г. № 1050-ст : введен впервые : дата введения 2019-07-01 / разработан Федеральным государственным унитарным предприятием "Информационное телеграфное агентство России (ИТАР-ТАСС)" филиал "Российская книжная палата", Федеральным государственным бюджетным учреждением "Российская государственная библиотека", Федеральным государственным бюджетным учреждением "Российская национальная библиотека". – Москва : Стандартинформ, 2018. – IV, 124, [1] c.
            
            formatted_text = text
            
        elif source_type == 'dissertation':
            # Форматирование для диссертаций и авторефератов
            # Фамилия, И. О. Название диссертации : специальность 00.00.00 «Название специальности» : диссертация на соискание ученой степени доктора / кандидата наук / И. О. Фамилия ; Организация. – Город, Год. – Число страниц с.
            
            formatted_text = text
            
        else:
            # Если тип не определен, оставляем как есть
            formatted_text = text
        
        # Добавляем номер обратно
        return f"{number_prefix}{formatted_text}" 

    def _correct_toc(self, document):
        """
        Исправляет оформление оглавления (содержания)
        ВАЖНО: НЕ использует paragraph.text = ... для сохранения форматирования
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            # Ищем раздел с оглавлением
            toc_started = False
            toc_paragraphs = []
            
            for i, paragraph in enumerate(document.paragraphs):
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip().lower()
                
                try:
                    # Определяем начало оглавления
                    if not toc_started and re.search(r'^(оглавление|содержание)$', text):
                        toc_started = True
                        # Форматируем заголовок оглавления
                        paragraph.style = document.styles['Heading 1']
                        pf = paragraph.paragraph_format
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        pf.first_line_indent = Cm(0)
                        pf.space_after = Pt(12)
                        pf.space_before = Pt(12)
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: убираем точку через runs
                        if paragraph.text.strip().endswith('.'):
                            if paragraph.runs:
                                last_run = paragraph.runs[-1]
                                last_run.text = last_run.text.rstrip('.')
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: upper через runs
                        for run in paragraph.runs:
                            run.text = run.text.upper()
                        continue
                    
                    # Собираем параграфы оглавления
                    if toc_started:
                        # Если встретили заголовок "ВВЕДЕНИЕ", значит оглавление закончилось
                        if text == 'введение' or paragraph.style.name.startswith('Heading'):
                            break
                        
                        # Добавляем параграф в список для дальнейшей обработки
                        toc_paragraphs.append((i, paragraph))
                
                except Exception as e:
                    print(f"ОШИБКА при поиске оглавления '{text[:50]}...': {str(e)}")
                    continue
            
            # Если нашли оглавление, форматируем его
            if toc_paragraphs:
                for i, paragraph in toc_paragraphs:
                    # Пропускаем пустые параграфы
                    if not paragraph.text.strip():
                        continue
                    
                    try:
                        # Форматируем элемент оглавления
                        pf = paragraph.paragraph_format
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        pf.first_line_indent = Cm(0)  # Без абзацного отступа
                        pf.left_indent = Cm(0)  # Без отступа слева
                        pf.space_after = Pt(0)
                        pf.space_before = Pt(0)
                        
                        # Применяем шрифт
                        for run in paragraph.runs:
                            if run.font.name != self.standard_rules['font']['name']:
                                run.font.name = self.standard_rules['font']['name']
                            if run.font.size != Pt(self.standard_rules['font']['size']):
                                run.font.size = Pt(self.standard_rules['font']['size'])
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: добавляем точки через runs
                        # Проверяем наличие точек между текстом и номером страницы
                        if not re.search(r'\.{2,}\s*\d+$', paragraph.text):
                            # Разделяем текст и номер страницы
                            match = re.search(r'^(.*?)\s*(\d+)$', paragraph.text)
                            if match and paragraph.runs:
                                text_part = match.group(1).strip()
                                page_number = match.group(2)
                                
                                # Создаём новый текст с точками
                                new_text = f"{text_part} {'.' * 20} {page_number}"
                                
                                # Очищаем все runs кроме первого
                                for run in paragraph.runs:
                                    run.text = ''
                                
                                # Устанавливаем новый текст в первый run
                                if paragraph.runs:
                                    paragraph.runs[0].text = new_text
                    
                    except Exception as e:
                        print(f"ОШИБКА при форматировании элемента оглавления '{paragraph.text[:50]}...': {str(e)}")
                        continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_toc: {str(e)}")
            import traceback
            traceback.print_exc()

    def _correct_appendices(self, document):
        """
        Исправляет оформление приложений
        ВАЖНО: НЕ использует paragraph.text = ... для сохранения форматирования
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            # Ищем начало раздела приложений
            appendix_started = False
            current_appendix = None
            
            for i, paragraph in enumerate(document.paragraphs):
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip()
                
                try:
                    # Проверяем, не является ли параграф началом приложения
                    if re.match(r'^ПРИЛОЖЕНИЕ\s+[А-Я]', text, re.IGNORECASE):
                        appendix_started = True
                        current_appendix = text
                        
                        # Форматируем заголовок приложения
                        paragraph.style = document.styles['Heading 1']
                        pf = paragraph.paragraph_format
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        pf.first_line_indent = Cm(0)
                        pf.space_after = Pt(12)
                        pf.space_before = Pt(12)
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: убираем точку через runs
                        if paragraph.text.strip().endswith('.'):
                            if paragraph.runs:
                                last_run = paragraph.runs[-1]
                                last_run.text = last_run.text.rstrip('.')
                        
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: upper через runs
                        for run in paragraph.runs:
                            run.text = run.text.upper()
                        
                        # Если после слова "ПРИЛОЖЕНИЕ" и буквы нет названия, ищем следующий параграф с названием
                        if len(paragraph.text.split()) < 3:
                            next_para_index = i + 1
                            if next_para_index < len(document.paragraphs):
                                next_para = document.paragraphs[next_para_index]
                                if id(next_para) not in table_paragraphs:
                                    if next_para.text.strip() and not next_para.style.name.startswith('Heading'):
                                        # Форматируем название приложения
                                        next_pf = next_para.paragraph_format
                                        next_pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        next_pf.first_line_indent = Cm(0)
                                        next_pf.space_after = Pt(12)
                                        next_pf.space_before = Pt(0)
                                        
                                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: upper через runs
                                        for run in next_para.runs:
                                            run.text = run.text.upper()
                    
                    # Форматируем содержимое приложения, если это не заголовок
                    elif appendix_started and not paragraph.style.name.startswith('Heading') and text:
                        # Для текста внутри приложения применяем стандартное форматирование
                        if not re.match(r'^(рисунок|рис\.|таблица)', text.lower()):
                            pf = paragraph.paragraph_format
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            pf.first_line_indent = Cm(1.25)
                            
                            # Восстанавливаем форматирование шрифта
                            for run in paragraph.runs:
                                if run.font.name != self.standard_rules['font']['name']:
                                    run.font.name = self.standard_rules['font']['name']
                                if run.font.size != Pt(self.standard_rules['font']['size']):
                                    run.font.size = Pt(self.standard_rules['font']['size'])
                        
                        # Для подписей к рисункам и таблицам применяем специальное форматирование
                        elif re.match(r'^(рисунок|рис\.)', text.lower()):
                            pf = paragraph.paragraph_format
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            pf.first_line_indent = Cm(0)
                        elif re.match(r'^таблица', text.lower()):
                            pf = paragraph.paragraph_format
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            pf.first_line_indent = Cm(0)
                
                except Exception as e:
                    print(f"ОШИБКА при обработке приложения '{text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_appendices: {str(e)}")
            import traceback
            traceback.print_exc()

    def _correct_text_accents(self, document):
        """
        Исправляет оформление акцентов в тексте (курсив, жирность)
        """
        # Проходим по всем параграфам
        for paragraph in document.paragraphs:
            # Пропускаем пустые параграфы
            if not paragraph.text.strip():
                continue
            
            # Пропускаем заголовки - у них свое форматирование
            if paragraph.style.name.startswith('Heading'):
                continue
            
            # Проверяем форматирование текста (акценты)
            has_inconsistent_formatting = False
            expected_bold = None
            expected_italic = None
            expected_font_size = None
            runs_count = len(paragraph.runs)
            
            # Пропускаем параграфы с одним run - они не требуют выравнивания форматирования
            if runs_count <= 1:
                continue
            
            # Если количество runs больше 1, проверяем необходимость выравнивания форматирования
            for i, run in enumerate(paragraph.runs):
                # Пропускаем пустые runs
                if not run.text.strip():
                    continue
                
                # Если это первый непустой run, запоминаем его форматирование как ожидаемое
                if expected_bold is None:
                    expected_bold = run.bold
                    expected_italic = run.italic
                    if hasattr(run.font, 'size') and run.font.size:
                        expected_font_size = run.font.size
                
                # Проверяем наличие несоответствий в форматировании
                if run.bold != expected_bold or run.italic != expected_italic:
                    # Если текст содержит всего несколько символов и отличается по форматированию,
                    # вероятно, это специально выделенный фрагмент (акцент)
                    if len(run.text.strip()) <= 5 or re.match(r'^[\s\.\,\:\;\"\'\(\)\[\]\-]+$', run.text):
                        # Сохраняем отдельное форматирование для акцентов
                        continue
                    
                    # Иначе считаем это ошибкой форматирования
                    has_inconsistent_formatting = True
                    break
                
                # Проверяем размер шрифта
                if hasattr(run.font, 'size') and run.font.size and expected_font_size and run.font.size != expected_font_size:
                    # Для небольших фрагментов текста допускаем разный размер шрифта
                    if len(run.text.strip()) <= 5:
                        continue
                    
                    has_inconsistent_formatting = True
                    break
            
            # Если обнаружено несогласованное форматирование, выравниваем его
            if has_inconsistent_formatting:
                # Используем основной стиль текста документа
                for run in paragraph.runs:
                    # Устанавливаем стандартный шрифт и размер
                    run.font.name = self.standard_rules['font']['name']
                    run.font.size = Pt(self.standard_rules['font']['size'])
                    
                    # Сбрасываем жирность и курсив
                    run.font.bold = False
                    run.font.italic = False

    def _correct_footnotes(self, document):
        """
        Исправляет оформление подстрочных ссылок
        """
        # Ищем все сноски в документе
        footnotes_found = False
        
        try:
            # Получаем доступ к сноскам документа через XML
            if hasattr(document, '_element') and hasattr(document._element, 'xpath'):
                # Найдем все элементы сносок в документе
                footnote_refs = document._element.xpath('//w:footnoteReference')
                
                if footnote_refs:
                    footnotes_found = True
                    # Сноски есть, проверяем их форматирование
                    
                    # Получаем доступ к части документа с содержимым сносок
                    footnotes_part = document._part.footnotes_part
                    
                    if footnotes_part:
                        footnotes_element = footnotes_part.element
                        
                        # Найдем все сноски
                        footnotes = footnotes_element.xpath('.//w:footnote')
                        
                        for footnote in footnotes:
                            # Проверяем, является ли это служебной сноской (разделитель)
                            footnote_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                            if footnote_id in ('0', '1'):  # Служебные сноски
                                continue
                            
                            # Находим параграфы внутри сноски
                            footnote_paras = footnote.xpath('.//w:p')
                            
                            for para in footnote_paras:
                                # Создаем объект Paragraph из XML элемента
                                p = Paragraph(para, document)
                                
                                # Форматируем текст сноски
                                p.paragraph_format.first_line_indent = Cm(0)
                                p.paragraph_format.left_indent = Cm(0)
                                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                p.paragraph_format.space_after = Pt(0)
                                p.paragraph_format.space_before = Pt(0)
                                
                                # Форматируем шрифт сноски (обычно шрифт для сносок - 10pt)
                                for run in p.runs:
                                    run.font.name = self.standard_rules['font']['name']
                                    run.font.size = Pt(10)  # Размер шрифта для сносок
                                    
                                    # Убираем курсив из URL (если это не ГОСТ)
                                    if 'http' in run.text:
                                        run.font.italic = False
            
            # Если сноски не были найдены через xpath, проверим наличие обычных сносок через API
            if not footnotes_found and hasattr(document, 'footnotes'):
                for footnote in document.footnotes:
                    # Форматируем параграфы сноски
                    for para in footnote.paragraphs:
                        para.paragraph_format.first_line_indent = Cm(0)
                        para.paragraph_format.left_indent = Cm(0)
                        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        
                        # Форматируем шрифт сноски
                        for run in para.runs:
                            run.font.name = self.standard_rules['font']['name']
                            run.font.size = Pt(10)  # Размер шрифта для сносок
                            
                            # Убираем курсив из URL
                            if 'http' in run.text:
                                run.font.italic = False
                
                footnotes_found = True
                        
        except Exception as e:
            print(f"Предупреждение: Не удалось исправить сноски: {str(e)}")
            # Продолжаем выполнение даже при ошибке, т.к. это некритичная функция
            
        return footnotes_found 

    def _correct_hyphenation(self, document):
        """
        Исправляет автоматические переносы в документе
        """
        for paragraph in document.paragraphs:
            # Пропускаем пустые параграфы и заголовки
            if not paragraph.text.strip() or paragraph.style.name.startswith('Heading'):
                continue
                
            # Включаем автоматические переносы для параграфа
            try:
                if paragraph._element.get_or_add_pPr():
                    # Создаем элемент для автоматической расстановки переносов
                    hyphenation_element = OxmlElement('w:suppressAutoHyphens')
                    hyphenation_element.set(qn('w:val'), '0')  # 0 = включено (не подавлять)
                    paragraph._element.get_or_add_pPr().append(hyphenation_element)
                    
                    # Добавляем настройку автоматического разрыва слов для русского языка
                    lang_element = OxmlElement('w:lang')
                    lang_element.set(qn('w:val'), 'ru-RU')
                    paragraph._element.get_or_add_pPr().append(lang_element)
            except Exception as e:
                print(f"Предупреждение: Не удалось настроить переносы: {str(e)}")
        
        # Исправляем неправильные переносы в тексте
        self._fix_incorrect_hyphenation(document)
        
        # Исправляем "висячие" предлоги и союзы
        self._fix_hanging_prepositions(document)

    def _fix_incorrect_hyphenation(self, document):
        """
        Исправляет неправильные переносы в тексте
        ВАЖНО: НЕ использует paragraph.text = ... для сохранения форматирования
        """
        try:
            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            # Список запрещенных переносов (слова, которые не должны переноситься)
            forbidden_hyphen_words = [
                r'\bи\b', r'\bа\b', r'\bв\b', r'\bс\b', r'\bк\b', r'\bу\b', r'\bо\b',
                r'\bна\b', r'\bот\b', r'\bдо\b', r'\bза\b', r'\bиз\b', r'\bпо\b',
                r'\bт\.д\b', r'\bт\.п\b', r'\bт\.е\b',
                r'\bг\.\b', r'\bгг\.\b', r'\bвв\.\b', r'\bстр\.\b'
            ]
            
            # Список правил для проверки и исправления неправильных переносов
            hyphen_rules = [
                (r'(\w)-\s+(\w)', r'\1\2'),  # Убираем переносы внутри слов
                (r'(\d+)\s*-\s*(\d+)', r'\1-\2'),  # Исправляем дефисы в числовых диапазонах
                (r'(\w+)\s*-\s*(\w+)', r'\1-\2')  # Исправляем дефисы между словами
            ]
            
            for paragraph in document.paragraphs:
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                # Пропускаем пустые параграфы
                if not paragraph.text.strip():
                    continue
                
                try:
                    # Проверяем необходимость изменений
                    needs_modification = False
                    for pattern in forbidden_hyphen_words:
                        if re.search(pattern, paragraph.text):
                            needs_modification = True
                            break
                    
                    if not needs_modification:
                        for pattern, _ in hyphen_rules:
                            if re.search(pattern, paragraph.text):
                                needs_modification = True
                                break
                    
                    if needs_modification and paragraph.runs:
                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: работаем через runs
                        for run in paragraph.runs:
                            if run.text:
                                text = run.text
                                
                                # Применяем замены для запрещенных переносов
                                for pattern in forbidden_hyphen_words:
                                    for match in re.finditer(r'\s+(' + pattern[2:-2] + r')\b', text):
                                        word = match.group(1)
                                        text = text.replace(f" {word}", f"\u00A0{word}")
                                
                                # Применяем правила для исправления переносов
                                for pattern, replacement in hyphen_rules:
                                    text = re.sub(pattern, replacement, text)
                                
                                run.text = text
                
                except Exception as e:
                    print(f"ОШИБКА при исправлении переносов '{paragraph.text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _fix_incorrect_hyphenation: {str(e)}")
            import traceback
            traceback.print_exc()

    def _fix_hanging_prepositions(self, document):
        """
        Исправляет "висячие" предлоги и союзы в конце строк, добавляя неразрывные пробелы
        """
        # Список предлогов и союзов, которые не должны находиться в конце строки
        hanging_words = [
            'а', 'и', 'в', 'с', 'к', 'у', 'о', 'на', 'от', 'до', 'за', 'из', 'по', 'под', 'над',
            'при', 'для', 'без', 'про', 'через', 'перед', 'после', 'кроме', 'вдоль', 'вместо',
            'около', 'возле', 'между', 'сквозь', 'среди', 'из-за', 'из-под', 'но', 'да', 'или',
            'либо', 'то', 'не', 'ни', 'бы', 'же', 'ведь', 'вот', 'что', 'как', 'так', 'уж'
        ]
        
        # Регулярное выражение для поиска предлогов и союзов в тексте
        # Ищем предлог/союз и пробел после него
        pattern = r'\b(' + '|'.join(hanging_words) + r')\s+'
        
        for paragraph in document.paragraphs:
            text = paragraph.text
            
            # Ищем все вхождения предлогов и союзов
            for match in re.finditer(pattern, text):
                # Получаем найденное слово и его позицию
                word = match.group(1)
                pos = match.start()
                
                # Заменяем обычный пробел после предлога на неразрывный
                text = text[:match.end()-1] + '\u00A0' + text[match.end():]
                
            # Если текст изменился, обновляем параграф
            if text != paragraph.text:
                paragraph.text = text
                
                # Восстанавливаем форматирование
                for run in paragraph.runs:
                    run.font.name = self.standard_rules['font']['name']
                    run.font.size = Pt(self.standard_rules['font']['size'])

    def _correct_cross_references(self, document):
        """
        Исправляет перекрестные ссылки в документе
        """
        # Словарь для хранения номеров рисунков, таблиц и формул
        reference_dict = {
            'рисунок': {},  # номер: заголовок
            'таблица': {},  # номер: заголовок
            'формула': {},  # номер: текст
            'раздел': {},   # номер: заголовок
            'приложение': {}  # буква: заголовок
        }
        
        # Первый проход - собираем информацию о номерах элементов
        for i, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            
            # Поиск рисунков
            if re.match(r'^рисунок\s+\d+', text.lower()) or re.match(r'^рис\.\s*\d+', text.lower()):
                match = re.search(r'(?:рисунок|рис\.)\s*(\d+)', text.lower())
                if match:
                    figure_num = match.group(1)
                    title = text[match.end():].strip()
                    if title.startswith('–') or title.startswith('-') or title.startswith('—'):
                        title = title[1:].strip()
                    reference_dict['рисунок'][figure_num] = title
            
            # Поиск таблиц
            elif re.match(r'^таблица\s+\d+', text.lower()):
                match = re.search(r'таблица\s+(\d+)', text.lower())
                if match:
                    table_num = match.group(1)
                    title = text[match.end():].strip()
                    if title.startswith('–') or title.startswith('-') or title.startswith('—'):
                        title = title[1:].strip()
                    reference_dict['таблица'][table_num] = title
            
            # Поиск формул с номерами
            elif '(' in text and ')' in text and len(text.strip()) < 50:  # Предположительно формула
                match = re.search(r'\((\d+(?:\.\d+)?)\)', text)
                if match:
                    formula_num = match.group(1)
                    reference_dict['формула'][formula_num] = text.replace(match.group(0), '').strip()
            
            # Поиск заголовков разделов
            elif paragraph.style.name.startswith('Heading'):
                match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)', text)
                if match:
                    section_num = match.group(1)
                    title = match.group(2)
                    reference_dict['раздел'][section_num] = title
            
            # Поиск приложений
            elif re.match(r'^приложение\s+[А-Я]', text.upper()):
                match = re.search(r'приложение\s+([А-Я])', text.upper())
                if match:
                    appendix_letter = match.group(1)
                    title = text[match.end():].strip()
                    reference_dict['приложение'][appendix_letter] = title
        
        # Второй проход - исправляем ссылки в тексте
        for paragraph in document.paragraphs:
            text = paragraph.text
            modified = False
            
            # Ищем ссылки на рисунки
            for match in re.finditer(r'(?<!\w)(рис\.|рисунк[а-я]*)\s*\.?\s*(\d+)(?!\d)', text, re.IGNORECASE):
                prefix, num = match.groups()
                if num in reference_dict['рисунок']:
                    correct_ref = f"рисунок {num}"
                    if text[match.start()-1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    modified = True
            
            # Ищем ссылки на таблицы
            for match in re.finditer(r'(?<!\w)(табл\.|таблиц[а-я]*)\s*\.?\s*(\d+)(?!\d)', text, re.IGNORECASE):
                prefix, num = match.groups()
                if num in reference_dict['таблица']:
                    correct_ref = f"таблица {num}"
                    if text[match.start()-1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    modified = True
            
            # Ищем ссылки на формулы
            for match in re.finditer(r'(?<!\w)(формул[а-я]*|выражени[а-я]*)\s*\.?\s*(\d+(?:\.\d+)?)(?!\d)', text, re.IGNORECASE):
                prefix, num = match.groups()
                if num in reference_dict['формула']:
                    correct_ref = f"формула ({num})"
                    if text[match.start()-1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    modified = True
            
            # Ищем ссылки на разделы
            for match in re.finditer(r'(?<!\w)(раздел[а-я]*|глав[а-я]*)\s*\.?\s*(\d+(?:\.\d+)?)(?!\d)', text, re.IGNORECASE):
                prefix, num = match.groups()
                if num in reference_dict['раздел']:
                    correct_ref = f"раздел {num}"
                    if text[match.start()-1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    modified = True
            
            # Ищем ссылки на приложения
            for match in re.finditer(r'(?<!\w)(приложени[а-я]*)\s*\.?\s*([А-Я])(?![А-Я])', text, re.IGNORECASE):
                prefix, letter = match.groups()
                if letter in reference_dict['приложение']:
                    correct_ref = f"приложение {letter}"
                    if text[match.start()-1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    modified = True
            
            # Обновляем текст параграфа, если были внесены изменения
            if modified:
                paragraph.text = text
                
                # Восстанавливаем форматирование после изменения текста
                for run in paragraph.runs:
                    run.font.name = self.standard_rules['font']['name']
                    run.font.size = Pt(self.standard_rules['font']['size'])

    def _correct_abbreviations_list(self, document):
        """
        Исправляет список сокращений и условных обозначений
        """
        # Ищем раздел со списком сокращений
        abbreviations_started = False
        abbreviations_paragraphs = []
        
        # Паттерны для идентификации раздела с сокращениями
        abbr_title_patterns = [
            r'список\s+сокращений',
            r'перечень\s+сокращений',
            r'список\s+условных\s+обозначений',
            r'условные\s+обозначения',
            r'принятые\s+сокращения'
        ]
        
        for i, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip().lower()
            
            # Определяем начало списка сокращений
            if not abbreviations_started:
                if any(re.search(pattern, text) for pattern in abbr_title_patterns):
                    abbreviations_started = True
                    # Форматируем заголовок списка сокращений
                    paragraph.style = document.styles['Heading 1']
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.space_after = Pt(12)
                    paragraph.paragraph_format.space_before = Pt(12)
                    
                    # Убираем точку в конце заголовка, если она есть
                    if paragraph.text.strip().endswith('.'):
                        paragraph.text = paragraph.text.strip().rstrip('.')
                    
                    # Приводим к верхнему регистру
                    paragraph.text = paragraph.text.upper()
                    continue
            
            # Собираем параграфы списка сокращений
            if abbreviations_started:
                # Если встретили новый заголовок, значит список сокращений закончился
                if paragraph.style.name.startswith('Heading'):
                    break
                
                # Добавляем параграф в список для дальнейшей обработки
                abbreviations_paragraphs.append((i, paragraph))
        
        # Если нашли список сокращений, форматируем его
        if abbreviations_paragraphs:
            # Создаем словарь сокращений для проверки и исправления
            abbreviations_dict = {}
            
            # Форматируем каждый элемент списка сокращений
            for i, paragraph in abbreviations_paragraphs:
                # Пропускаем пустые параграфы
                if not paragraph.text.strip():
                    continue
                
                text = paragraph.text.strip()
                
                # Определяем формат элемента списка сокращений (обычно это "Сокращение – расшифровка")
                parts = re.split(r'\s+[-–—]\s+', text, 1)
                
                if len(parts) == 2:
                    abbr, description = parts
                    abbreviations_dict[abbr.strip()] = description.strip()
                    
                    # Форматируем элемент списка
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.left_indent = Cm(0)
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    
                    # Проверяем, есть ли точка в конце расшифровки
                    if not description.strip().endswith('.') and not description.strip().endswith(':'):
                        # Добавляем точку в конце
                        paragraph.text = f"{abbr.strip()} – {description.strip()}."
                
                # Восстанавливаем форматирование после изменения текста
                for run in paragraph.runs:
                    run.font.name = self.standard_rules['font']['name']
                    run.font.size = Pt(self.standard_rules['font']['size'])
            
            # Проверяем использование сокращений в тексте
            self._check_abbreviations_usage(document, abbreviations_dict)
            
            return True
        
        return False
    
    def _check_abbreviations_usage(self, document, abbreviations_dict):
        """
        Проверяет правильность использования сокращений в тексте
        
        Args:
            document: Документ Word
            abbreviations_dict: Словарь сокращений (сокращение: расшифровка)
        """
        if not abbreviations_dict:
            return
        
        # Находим первое использование каждого сокращения и проверяем, 
        # было ли оно расшифровано при первом употреблении
        abbreviation_first_use = {}
        
        for i, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            
            # Ищем сокращения в тексте
            for abbr in abbreviations_dict.keys():
                # Используем регулярное выражение для поиска сокращения как отдельного слова
                pattern = r'\b' + re.escape(abbr) + r'\b'
                match = re.search(pattern, text)
                
                if match and abbr not in abbreviation_first_use:
                    # Сохраняем информацию о первом использовании
                    abbreviation_first_use[abbr] = (i, paragraph)
                    
                    # Проверяем, есть ли расшифровка в этом же параграфе
                    description = abbreviations_dict[abbr]
                    if description.lower() not in text.lower() and f"({abbr})" not in text:
                        # Если расшифровки нет, добавляем ее в скобках после первого употребления
                        text_before = text[:match.end()]
                        text_after = text[match.end():]
                        new_text = f"{text_before} ({description}){text_after}"
                        paragraph.text = new_text