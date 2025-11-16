from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE

# Создание документа

doc = Document()

# Настройка полей
section = doc.sections[0]
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# Стили
styles = doc.styles

# Основной текст
style = styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(1.25)

# Заголовок главы
if 'Heading 1' in styles:
    heading1 = styles['Heading 1']
else:
    heading1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
heading1.font.name = 'Times New Roman'
heading1.font.size = Pt(14)
heading1.font.bold = True
heading1.font.color.rgb = RGBColor(0, 0, 0)
heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading1.paragraph_format.space_after = Pt(12)
heading1.paragraph_format.first_line_indent = 0

# Заголовок подраздела
if 'Heading 2' in styles:
    heading2 = styles['Heading 2']
else:
    heading2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
heading2.font.name = 'Times New Roman'
heading2.font.size = Pt(14)
heading2.font.bold = True
heading2.font.color.rgb = RGBColor(0, 0, 0)
heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
heading2.paragraph_format.first_line_indent = Cm(1.25)
heading2.paragraph_format.space_after = Pt(12)

# Стиль для подписей к рисункам и таблицам
if 'Caption' in styles:
    caption = styles['Caption']
else:
    caption = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
caption.font.name = 'Times New Roman'
caption.font.size = Pt(12)
caption.font.italic = True
caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.first_line_indent = 0

# Титульный лист (пример)
doc.add_paragraph('МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РФ', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('ФГБОУ ВО «ВАШ ВУЗ»', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('ИНСТИТУТ/ФАКУЛЬТЕТ', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('КАФЕДРА', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Тема: «Название темы»', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Студент: Иванов И.И.', style='Normal').alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph('Руководитель: Петров П.П.', style='Normal').alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph('Город, 2024', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Оглавление (автоматическое поле)
toc = doc.add_paragraph('ОГЛАВЛЕНИЕ', style='Heading 1')
# Добавим поле оглавления (TOC)
run = toc.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar)
instrText = OxmlElement('w:instrText')
instrText.text = 'TOC \\o \\u'
run._r.append(instrText)
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'separate')
run._r.append(fldChar)
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar)

doc.add_page_break()

# Введение
doc.add_paragraph('ВВЕДЕНИЕ', style='Heading 1')
doc.add_paragraph('Текст введения...')

doc.add_page_break()

# Пример главы и подраздела
doc.add_paragraph('1 ХАРАКТЕРИСТИКА ДЕТАЛИ «ОСЬ»', style='Heading 1')
doc.add_paragraph('1.1 Химический состав сплавов', style='Heading 2')
doc.add_paragraph('Основной текст главы...')

# Пример перечисления
p = doc.add_paragraph('Курсы обучения:', style='Normal')
p.paragraph_format.first_line_indent = Cm(1.25)
doc.add_paragraph('– первый,', style='Normal')
doc.add_paragraph('– второй,', style='Normal')
doc.add_paragraph('– третий.', style='Normal')

# Пример сложного перечисления
p = doc.add_paragraph('Курсы обучения:', style='Normal')
p.paragraph_format.first_line_indent = Cm(1.25)
doc.add_paragraph('– первый:', style='Normal')
doc.add_paragraph('   а) математика-физика;', style='Normal')
doc.add_paragraph('   б) информатика-математика;', style='Normal')
doc.add_paragraph('– второй:', style='Normal')
doc.add_paragraph('   а) информатика-физика;', style='Normal')
doc.add_paragraph('   б) информатика-математика.', style='Normal')

# Пример таблицы
table = doc.add_table(rows=2, cols=3)
table.style = 'Table Grid'
table.autofit = False
table.allow_autofit = False
for cell in table.rows[0].cells:
    cell.width = Cm(4)
table.cell(0, 0).text = 'Фамилия'
table.cell(0, 1).text = 'Имя'
table.cell(0, 2).text = 'Возраст, г'
table.cell(1, 0).text = 'Иванов'
table.cell(1, 1).text = 'Петр'
table.cell(1, 2).text = '15'
doc.add_paragraph('Таблица 1 – Структурные элементы', style='Caption')

# Пример рисунка (заглушка)
doc.add_paragraph('Рисунок 1 – Структурные элементы', style='Caption')

# Пример ссылки
p = doc.add_paragraph('Итоги представлены в таблице 1. В соответствии с рисунком 2. Рассмотрим формулу (1).', style='Normal')

# Пример списка литературы
doc.add_page_break()
doc.add_paragraph('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', style='Heading 1')
doc.add_paragraph('1. Каменский, П. П. Труды по истории изобразительного искусства...', style='Normal')
doc.add_paragraph('2. Новиков, Н. И. Опыт исторического словаря...', style='Normal')

# Пример приложения
doc.add_page_break()
doc.add_paragraph('ПРИЛОЖЕНИЯ', style='Heading 1')
doc.add_page_break()
doc.add_paragraph('Приложение А', style='Heading 1')
doc.add_paragraph('Наименование приложения', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Сохраняем шаблон
doc.save('backend/app/templates/norm_control_template.docx')
print('Шаблон norm_control_template.docx успешно создан!') 