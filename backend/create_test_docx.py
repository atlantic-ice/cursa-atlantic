import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_test_document(output_path='test_document.docx'):
    """Создает тестовый документ с различными элементами для проверки обработки"""
    doc = Document()
    
    # Установка полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Добавляем заголовок
    title = doc.add_heading('Тестовый документ для проверки', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Добавляем обычный текст
    doc.add_paragraph('Это обычный параграф для тестирования обработки текста.')
    
    # Добавляем форматированный текст
    p = doc.add_paragraph('Этот параграф содержит ')
    p.add_run('жирный текст').bold = True
    p.add_run(' и ')
    p.add_run('курсив').italic = True
    p.add_run('.')
    
    # Добавляем список
    doc.add_paragraph('Ниже следует маркированный список:', style='List Bullet')
    doc.add_paragraph('Первый пункт списка', style='List Bullet')
    doc.add_paragraph('Второй пункт списка', style='List Bullet')
    doc.add_paragraph('Третий пункт списка', style='List Bullet')
    
    # Добавляем нумерованный список
    doc.add_paragraph('Ниже следует нумерованный список:', style='List Number')
    doc.add_paragraph('Первый пункт', style='List Number')
    doc.add_paragraph('Второй пункт', style='List Number')
    doc.add_paragraph('Третий пункт', style='List Number')
    
    # Добавляем таблицу
    doc.add_paragraph('Пример таблицы:')
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Заполняем заголовок таблицы
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Заголовок 1'
    heading_cells[1].text = 'Заголовок 2'
    heading_cells[2].text = 'Заголовок 3'
    
    # Заполняем таблицу данными
    for i in range(1, 3):
        row_cells = table.rows[i].cells
        for j in range(3):
            row_cells[j].text = f'Ячейка {i},{j+1}'
    
    # Добавляем подзаголовок
    doc.add_heading('Подраздел документа', level=2)
    doc.add_paragraph('Текст в подразделе документа для проверки обработки заголовков.')
    
    # Добавляем еще один подраздел
    doc.add_heading('Еще один подраздел', level=2)
    doc.add_paragraph('Дополнительный текст для тестирования.')
    
    # Добавляем информацию о библиографии
    doc.add_heading('Список литературы', level=1)
    doc.add_paragraph('1. Иванов И.И. Название книги. - М.: Издательство, 2023. - 123 с.')
    doc.add_paragraph('2. Петров П.П. Другая книга // Журнал. - 2024. - №5. - С. 45-67.')
    doc.add_paragraph('3. Сидоров С.С. Электронный ресурс [Электронный ресурс]. - URL: https://example.com (дата обращения: 01.01.2024).')
    
    # Сохраняем документ
    doc.save(output_path)
    print(f'Документ успешно создан: {os.path.abspath(output_path)}')
    return os.path.abspath(output_path)

if __name__ == "__main__":
    create_test_document() 