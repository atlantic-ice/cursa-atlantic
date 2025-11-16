"""
Тестирование процесса исправления документа
"""
from docx import Document
from app.services.document_corrector import DocumentCorrector
import os

def test_correction():
    """Тестирует исправление документа"""
    
    # Путь к тестовому документу
    input_path = 'test_data/test_document_with_errors.docx'
    output_path = 'test_data/test_document_corrected.docx'
    
    print(f"Загрузка документа: {input_path}")
    doc = Document(input_path)
    
    # Подсчитываем элементы ДО исправления
    print("\n=== ДО ИСПРАВЛЕНИЯ ===")
    print(f"Параграфов: {len(doc.paragraphs)}")
    print(f"Таблиц: {len(doc.tables)}")
    print(f"Секций: {len(doc.sections)}")
    
    # Выводим содержимое первых параграфов
    print("\nПервые 5 параграфов:")
    for i, p in enumerate(doc.paragraphs[:5]):
        if p.text.strip():
            print(f"  {i+1}. '{p.text[:60]}...' (шрифт: {p.runs[0].font.name if p.runs else 'N/A'}, размер: {p.runs[0].font.size if p.runs else 'N/A'})")
    
    # Создаем корректор и исправляем документ
    print("\n=== ПРОЦЕСС ИСПРАВЛЕНИЯ ===")
    corrector = DocumentCorrector()
    
    # Сохраняем документ во временный файл для обработки
    temp_input = 'test_data/temp_input.docx'
    doc.save(temp_input)
    
    # Исправляем все ошибки
    corrected_path = corrector.correct_document(temp_input, out_path=output_path, errors=None)
    
    # Загружаем исправленный документ
    result_doc = Document(corrected_path)
    
    # Подсчитываем элементы ПОСЛЕ исправления
    print("\n=== ПОСЛЕ ИСПРАВЛЕНИЯ ===")
    print(f"Параграфов: {len(result_doc.paragraphs)}")
    print(f"Таблиц: {len(result_doc.tables)}")
    print(f"Секций: {len(result_doc.sections)}")
    
    # Выводим содержимое первых параграфов после исправления
    print("\nПервые 5 параграфов:")
    for i, p in enumerate(result_doc.paragraphs[:5]):
        if p.text.strip():
            font_name = p.runs[0].font.name if p.runs else 'N/A'
            font_size = p.runs[0].font.size if p.runs else 'N/A'
            alignment = p.alignment
            first_line = p.paragraph_format.first_line_indent
            print(f"  {i+1}. '{p.text[:60]}...'")
            print(f"      Шрифт: {font_name}, Размер: {font_size}")
            print(f"      Выравнивание: {alignment}, Отступ: {first_line}")
    
    # Проверяем поля
    print("\nПараметры страницы:")
    section = result_doc.sections[0]
    print(f"  Левое поле: {section.left_margin.cm:.2f} см (должно быть 3.0)")
    print(f"  Правое поле: {section.right_margin.cm:.2f} см (должно быть 1.0)")
    print(f"  Верхнее поле: {section.top_margin.cm:.2f} см (должно быть 2.0)")
    print(f"  Нижнее поле: {section.bottom_margin.cm:.2f} см (должно быть 2.0)")
    
    # Проверяем, что файл создан и не пустой
    if os.path.exists(output_path):
        file_size = os.path.getsize(output_path)
        print(f"\n✓ Исправленный документ сохранен: {output_path}")
        print(f"  Размер файла: {file_size:,} байт")
        
        # Открываем сохраненный файл для проверки
        saved_doc = Document(output_path)
        print(f"  Параграфов в сохраненном файле: {len(saved_doc.paragraphs)}")
        print(f"  Таблиц в сохраненном файле: {len(saved_doc.tables)}")
        
        # Проверяем, что содержимое не потеряно
        if len(saved_doc.paragraphs) == 0:
            print("\n❌ ОШИБКА: Весь контент был удален!")
        elif len(saved_doc.paragraphs) < len(doc.paragraphs) - 5:
            print(f"\n⚠️ ВНИМАНИЕ: Потеряно {len(doc.paragraphs) - len(saved_doc.paragraphs)} параграфов!")
        else:
            print("\n✓ Контент сохранен!")
    else:
        print(f"\n❌ ОШИБКА: Файл не был создан!")
    
    # Удаляем временный файл
    if os.path.exists(temp_input):
        os.remove(temp_input)

if __name__ == '__main__':
    test_correction()
