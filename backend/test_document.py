"""
Создание тестового документа с нарушениями ГОСТ для проверки исправлений
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_test_document():
    """Создает тестовый документ с нарушениями ГОСТ"""
    doc = Document()
    
    # Неправильные поля (должны быть: левое 3см, правое 1см, верхнее 2см, нижнее 2см)
    section = doc.sections[0]
    section.left_margin = Cm(2.0)  # НЕПРАВИЛЬНО - должно быть 3
    section.right_margin = Cm(2.0)  # НЕПРАВИЛЬНО - должно быть 1
    section.top_margin = Cm(1.5)   # НЕПРАВИЛЬНО - должно быть 2
    section.bottom_margin = Cm(1.5) # НЕПРАВИЛЬНО - должно быть 2
    
    # Заголовок 1 уровня с нарушениями
    h1 = doc.add_paragraph("1. введение")  # должно быть заглавными
    h1.style = 'Heading 1'
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # НЕПРАВИЛЬНО - должно быть CENTER
    for run in h1.runs:
        run.font.name = 'Arial'  # НЕПРАВИЛЬНО - должно быть Times New Roman
        run.font.size = Pt(14)   # НЕПРАВИЛЬНО - должно быть 16
    
    # Основной текст с нарушениями
    p1 = doc.add_paragraph(
        "Это первый абзац основного текста документа. "
        "Текст не имеет абзацного отступа и неправильно выровнен."
    )
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # НЕПРАВИЛЬНО - должно быть JUSTIFY
    for run in p1.runs:
        run.font.name = 'Arial'        # НЕПРАВИЛЬНО - должно быть Times New Roman
        run.font.size = Pt(12)         # НЕПРАВИЛЬНО - должно быть 14
    p1.paragraph_format.first_line_indent = Cm(0)  # НЕПРАВИЛЬНО - должно быть 1.25
    p1.paragraph_format.line_spacing = 1.0         # НЕПРАВИЛЬНО - должно быть 1.5
    
    # Второй абзац
    p2 = doc.add_paragraph(
        "Второй абзац текста также имеет нарушения форматирования. "
        "Межстрочный интервал одинарный вместо полуторного."
    )
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in p2.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    p2.paragraph_format.first_line_indent = Cm(0.5)  # НЕПРАВИЛЬНО
    p2.paragraph_format.line_spacing = 1.0           # НЕПРАВИЛЬНО
    
    # Заголовок 2 уровня с нарушениями
    h2 = doc.add_paragraph("1.1 Актуальность темы")
    h2.style = 'Heading 2'
    h2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # НЕПРАВИЛЬНО - должно быть LEFT
    for run in h2.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)  # НЕПРАВИЛЬНО - должно быть 14
        run.font.bold = False   # НЕПРАВИЛЬНО - должно быть True
    
    # Текст после заголовка
    p3 = doc.add_paragraph(
        "Текст после заголовка второго уровня. "
        "Этот текст также содержит множество нарушений ГОСТ."
    )
    for run in p3.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)  # НЕПРАВИЛЬНО - должно быть 14
    p3.paragraph_format.first_line_indent = Cm(2.0)  # НЕПРАВИЛЬНО - должно быть 1.25
    
    # Неправильная подпись к рисунку
    img_caption = doc.add_paragraph("рис. 1 - График функции")  # должно быть "Рисунок 1 – График функции."
    img_caption.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # НЕПРАВИЛЬНО - должно быть CENTER
    for run in img_caption.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    
    # Таблица с нарушениями
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Заполняем заголовок таблицы
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Значение'
    header_cells[2].text = 'Примечание'
    
    # Заполняем данные с неправильным форматированием
    for i in range(1, 3):
        cells = table.rows[i].cells
        cells[0].text = f'Параметр {i}'
        cells[1].text = f'Значение {i}'
        cells[2].text = f'Примечание {i}'
        
        # Неправильный шрифт в ячейках
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'  # НЕПРАВИЛЬНО
                    run.font.size = Pt(10)   # НЕПРАВИЛЬНО
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # НЕПРАВИЛЬНО
    
    # Неправильное название таблицы
    table_caption = doc.add_paragraph("Таблица 1 - Параметры системы")  # должна быть точка в конце
    table_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # НЕПРАВИЛЬНО - должно быть LEFT
    
    # Маркированный список с нарушениями
    doc.add_paragraph("Основные преимущества:")
    list_item1 = doc.add_paragraph("• высокая производительность")
    list_item1.paragraph_format.first_line_indent = Cm(0)  # НЕПРАВИЛЬНО
    for run in list_item1.runs:
        run.font.size = Pt(12)  # НЕПРАВИЛЬНО
    
    list_item2 = doc.add_paragraph("• низкая стоимость")
    list_item2.paragraph_format.first_line_indent = Cm(0)  # НЕПРАВИЛЬНО
    
    list_item3 = doc.add_paragraph("• простота использования")
    list_item3.paragraph_format.first_line_indent = Cm(0)  # НЕПРАВИЛЬНО
    
    # Нумерованный список с нарушениями
    doc.add_paragraph("Этапы разработки:")
    num_item1 = doc.add_paragraph("1) анализ требований")
    num_item1.paragraph_format.first_line_indent = Cm(0)  # НЕПРАВИЛЬНО
    
    num_item2 = doc.add_paragraph("2) проектирование")
    num_item2.paragraph_format.first_line_indent = Cm(0)  # НЕПРАВИЛЬНО
    
    num_item3 = doc.add_paragraph("3) реализация")
    num_item3.paragraph_format.first_line_indent = Cm(0)  # НЕПРАВИЛЬНО
    
    # Сохраняем документ
    output_path = 'test_data/test_document_with_errors.docx'
    doc.save(output_path)
    print(f"✓ Тестовый документ создан: {output_path}")
    print("\nСодержит следующие нарушения ГОСТ:")
    print("- Неправильные поля страницы")
    print("- Неправильный шрифт (Arial, Calibri вместо Times New Roman)")
    print("- Неправильный размер шрифта")
    print("- Неправильное выравнивание текста")
    print("- Отсутствие/неправильный абзацный отступ")
    print("- Одинарный межстрочный интервал вместо полуторного")
    print("- Неправильное форматирование заголовков")
    print("- Неправильные подписи к рисункам и таблицам")
    print("- Неправильное форматирование списков")
    
    return output_path

if __name__ == '__main__':
    create_test_document()
