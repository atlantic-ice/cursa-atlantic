import os
import sys
import tempfile
from docx import Document

def test_docx_create():
    """Проверяет возможность создания и сохранения docx файла"""
    print("Проверка создания DOCX файла...")
    
    # Создаем новый документ
    doc = Document()
    doc.add_paragraph("Тестовый документ")
    
    # Сохраняем во временный файл
    temp_dir = tempfile.mkdtemp()
    test_file = os.path.join(temp_dir, "test_doc.docx")
    print(f"Сохранение в: {test_file}")
    doc.save(test_file)
    
    # Проверяем, что файл создан
    if os.path.exists(test_file):
        print(f"Успешно: файл создан, размер: {os.path.getsize(test_file)} байт")
    else:
        print("Ошибка: файл не создан")
        return False
    
    # Пробуем открыть сохраненный файл
    try:
        reopened_doc = Document(test_file)
        print(f"Успешно: файл открыт повторно, {len(reopened_doc.paragraphs)} параграфов")
    except Exception as e:
        print(f"Ошибка при открытии файла: {type(e).__name__}: {str(e)}")
        return False
    
    # Очищаем временный файл
    os.unlink(test_file)
    os.rmdir(temp_dir)
    return True

def print_environment_info():
    """Выводит информацию о Python и окружении"""
    print(f"Python: {sys.version}")
    print(f"Текущая директория: {os.getcwd()}")
    
    try:
        import docx
        print(f"python-docx версия: {docx.__version__}")
    except Exception as e:
        print(f"Ошибка при импорте docx: {str(e)}")
    
    # Проверяем наличие других связанных библиотек
    try:
        import lxml
        print(f"lxml версия: {lxml.__version__}")
    except Exception as e:
        print(f"lxml не установлен или возникла ошибка: {str(e)}")

if __name__ == "__main__":
    print("=== Проверка библиотеки python-docx ===")
    print_environment_info()
    print("\n=== Тест создания документа ===")
    result = test_docx_create()
    print("\nИтоговый результат:", "УСПЕШНО" if result else "ОШИБКА") 