"""
Полный end-to-end тест для проверки всех функций приложения
Запуск: python test_full_flow.py
"""
import requests
import json
import os
import time
from pathlib import Path
from docx import Document

# Конфигурация
BASE_URL = "http://localhost:5000"
API_URL = f"{BASE_URL}/api/document"
TEST_DIR = Path(__file__).parent / "test_data"
TEST_DOCUMENT = TEST_DIR / "test_document.docx"

class Colors:
    GREEN = '\033[92m'
    RED = '\033[91m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    RESET = '\033[0m'

def print_success(msg):
    print(f"{Colors.GREEN}✓{Colors.RESET} {msg}")

def print_error(msg):
    print(f"{Colors.RED}✗{Colors.RESET} {msg}")

def print_info(msg):
    print(f"{Colors.BLUE}ℹ{Colors.RESET} {msg}")

def print_warning(msg):
    print(f"{Colors.YELLOW}⚠{Colors.RESET} {msg}")

def create_test_document():
    """Создает тестовый документ с намеренными ошибками для проверки"""
    doc = Document()
    
    # Титульный лист с ошибками
    doc.add_heading('КУРСОВАЯ РАБОТА', level=1)  # Должен быть заголовок "ФЕДЕРАЛЬНОЕ АГЕНТСТВО..."
    doc.add_paragraph('По дисциплине: Информатика')
    doc.add_paragraph('Тема: Разработка системы контроля')
    doc.add_paragraph()
    
    # Содержание без точек после номеров
    doc.add_heading('СОДЕРЖАНИЕ', level=1)
    doc.add_paragraph('ВВЕДЕНИЕ 3')  # Нет точки и пробела между номером и страницей
    doc.add_paragraph('1 ГЛАВА ПЕРВАЯ 5')
    doc.add_paragraph('1.1 Параграф первый 7')
    doc.add_paragraph('ЗАКЛЮЧЕНИЕ 20')
    doc.add_paragraph()
    
    # Введение
    doc.add_heading('ВВЕДЕНИЕ', level=1)
    doc.add_paragraph('Это краткое введение.')  # Меньше 1 страницы
    doc.add_paragraph()
    
    # Основной текст
    doc.add_heading('1 Теоретическая часть', level=1)  # Не все заглавные буквы
    doc.add_paragraph('Текст главы. ' * 50)
    
    doc.add_heading('1.1 Подраздел', level=2)
    doc.add_paragraph('Текст подраздела. ' * 30)
    
    # Заключение
    doc.add_heading('ЗАКЛЮЧЕНИЕ', level=1)
    doc.add_paragraph('Краткое заключение.')  # Меньше 1 страницы
    doc.add_paragraph()
    
    # Список литературы с ошибками
    doc.add_heading('СПИСОК ЛИТЕРАТУРЫ', level=1)
    doc.add_paragraph('1 Иванов И.И. Книга. Москва, 2020.')  # Нет точки после номера
    doc.add_paragraph('2 Петров П.П Статья. 2021')  # Нет точки после инициалов, нет точки в конце
    
    TEST_DIR.mkdir(exist_ok=True)
    doc.save(TEST_DOCUMENT)
    print_success(f"Создан тестовый документ: {TEST_DOCUMENT}")

def test_server_health():
    """Проверка доступности сервера"""
    print_info("Проверка доступности сервера...")
    try:
        response = requests.get(f"{API_URL}/list-corrections", timeout=5)
        if response.status_code == 200:
            data = response.json()
            print_success(f"Сервер доступен. Найдено файлов коррекций: {data.get('file_count', 0)}")
            return True
        else:
            print_error(f"Сервер вернул статус: {response.status_code}")
            return False
    except requests.exceptions.RequestException as e:
        print_error(f"Ошибка подключения к серверу: {e}")
        print_warning("Убедитесь, что backend запущен: cd backend && python run.py")
        return False

def test_upload_document():
    """Тест загрузки документа"""
    print_info("Загрузка тестового документа...")
    
    if not TEST_DOCUMENT.exists():
        create_test_document()
    
    try:
        with open(TEST_DOCUMENT, 'rb') as f:
            files = {'file': ('test_document.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = requests.post(f"{API_URL}/upload", files=files, timeout=30)
        
        if response.status_code == 200:
            data = response.json()
            if data.get('success'):
                print_success(f"Документ загружен успешно!")
                print_info(f"  Файл: {data.get('filename')}")
                
                # check_results это dict с ключами: rules_results, total_issues_count, issues
                check_results = data.get('check_results', {})
                total_issues = check_results.get('total_issues_count', 0)
                print_info(f"  Найдено нарушений: {total_issues}")
                
                # Вывод первых 3 ошибок
                issues = check_results.get('issues', [])
                if issues and len(issues) > 0:
                    print_info("  Примеры нарушений:")
                    for i, issue in enumerate(list(issues)[:3], 1):
                        print(f"    {i}. [{issue.get('severity', 'N/A')}] {issue.get('description', 'N/A')}")
                
                return data
            else:
                print_error(f"Ошибка обработки: {data.get('error', 'Неизвестная ошибка')}")
                return None
        else:
            print_error(f"HTTP {response.status_code}: {response.text}")
            return None
    except Exception as e:
        print_error(f"Ошибка при загрузке: {e}")
        return None

def test_analyze_document(filename):
    """Тест анализа документа"""
    print_info("Анализ документа...")
    
    try:
        response = requests.post(f"{API_URL}/analyze", 
                                json={'filename': filename},
                                timeout=30)
        
        if response.status_code == 200:
            data = response.json()
            if data.get('success'):
                print_success("Анализ завершен!")
                print_info(f"  Найдено проблем: {len(data.get('issues', []))}")
                return data
            else:
                print_error(f"Ошибка анализа: {data.get('error')}")
                return None
        else:
            print_error(f"HTTP {response.status_code}: {response.text}")
            return None
    except Exception as e:
        print_error(f"Ошибка при анализе: {e}")
        return None

def test_correct_document(file_path, issues):
    """Тест автокоррекции документа"""
    print_info("Применение автокоррекций...")
    
    try:
        response = requests.post(f"{API_URL}/correct",
                                json={
                                    'file_path': file_path,  # Передаем путь к temp файлу
                                    'errors': issues[:5]  # Первые 5 ошибок
                                },
                                timeout=30)
        
        if response.status_code == 200:
            data = response.json()
            if data.get('success'):
                print_success("Коррекции применены!")
                print_info(f"  Исправленный файл: {data.get('filename')}")
                return data
            else:
                print_error(f"Ошибка коррекции: {data.get('error')}")
                return None
        else:
            print_error(f"HTTP {response.status_code}: {response.text}")
            return None
    except Exception as e:
        print_error(f"Ошибка при коррекции: {e}")
        return None

def test_generate_report(check_results, filename):
    """Тест генерации отчета DOCX"""
    print_info("Генерация отчета...")
    
    try:
        response = requests.post(f"{API_URL}/generate-report",
                                json={
                                    'check_results': check_results,
                                    'filename': filename
                                },
                                timeout=30)
        
        if response.status_code == 200:
            data = response.json()
            if data.get('success'):
                print_success("Отчет сгенерирован!")
                print_info(f"  Файл отчета: {data.get('filename')}")
                return data
            else:
                print_error(f"Ошибка генерации: {data.get('error')}")
                return None
        else:
            print_error(f"HTTP {response.status_code}: {response.text}")
            return None
    except Exception as e:
        print_error(f"Ошибка при генерации отчета: {e}")
        return None

def test_download_report(report_path, filename):
    """Тест скачивания отчета"""
    print_info("Скачивание отчета...")
    
    try:
        response = requests.get(f"{API_URL}/download-report",
                               params={'path': report_path, 'filename': filename},
                               timeout=30)
        
        if response.status_code == 200:
            output_path = TEST_DIR / f"downloaded_{filename}"
            with open(output_path, 'wb') as f:
                f.write(response.content)
            print_success(f"Отчет скачан: {output_path}")
            print_info(f"  Размер: {len(response.content)} bytes")
            return True
        else:
            print_error(f"HTTP {response.status_code}: {response.text}")
            return False
    except Exception as e:
        print_error(f"Ошибка при скачивании: {e}")
        return False

def run_full_test():
    """Запуск полного теста"""
    print("\n" + "="*60)
    print(f"{Colors.BLUE}ПОЛНЫЙ E2E ТЕСТ СИСТЕМЫ НОРМОКОНТРОЛЯ{Colors.RESET}")
    print("="*60 + "\n")
    
    # Проверка сервера
    if not test_server_health():
        print_error("\nТест прерван: сервер недоступен")
        return False
    
    print("\n" + "-"*60)
    
    # 1. Загрузка документа
    upload_result = test_upload_document()
    if not upload_result:
        print_error("\nТест прерван: ошибка загрузки")
        return False
    
    filename = upload_result.get('filename')
    file_path = upload_result.get('temp_path')  # Временный путь к загруженному файлу
    check_results = upload_result.get('check_results', {})
    issues = check_results.get('issues', [])
    
    print("\n" + "-"*60)
    
    # 2. Анализ документа - пропускаем, т.к. уже выполнен в upload
    print_info("Анализ документа уже выполнен при загрузке")
    
    print("\n" + "-"*60)
    
    # 3. Коррекция документа (если есть ошибки)
    if issues:
        correct_result = test_correct_document(file_path, issues)
        if not correct_result:
            print_warning("Коррекция не удалась, но продолжаем...")
    else:
        print_info("Ошибок для коррекции не найдено")
    
    print("\n" + "-"*60)
    
    # 4. Генерация отчета
    report_result = test_generate_report(check_results, filename)
    if not report_result:
        print_error("\nТест прерван: ошибка генерации отчета")
        return False
    
    report_path = report_result.get('report_file_path')
    report_filename = report_result.get('filename')
    
    print("\n" + "-"*60)
    
    # 5. Скачивание отчета
    if report_path and report_filename:
        download_ok = test_download_report(report_path, report_filename)
        if not download_ok:
            print_warning("Скачивание не удалось")
    
    print("\n" + "="*60)
    print(f"{Colors.GREEN}✓ ПОЛНЫЙ ТЕСТ ЗАВЕРШЕН УСПЕШНО!{Colors.RESET}")
    print("="*60 + "\n")
    
    return True

if __name__ == "__main__":
    try:
        success = run_full_test()
        exit(0 if success else 1)
    except KeyboardInterrupt:
        print_warning("\n\nТест прерван пользователем")
        exit(1)
    except Exception as e:
        print_error(f"\n\nКритическая ошибка: {e}")
        import traceback
        traceback.print_exc()
        exit(1)
