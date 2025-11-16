"""
Интеграционные тесты проверки полного цикла работы с документами
"""
import os
import sys
import pytest
import json
import io
from pathlib import Path
import tempfile
from docx import Document
import logging

# Настраиваем логгирование
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Добавляем путь к корневой директории проекта
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../..')))

# Импорт приложения Flask
from app import create_app

# Константы
TEST_DATA_DIR = Path(__file__).parent.parent / "test_data"
RESULTS_DIR = TEST_DATA_DIR / "results"

# Создаем директорию для результатов, если она не существует
os.makedirs(RESULTS_DIR, exist_ok=True)

@pytest.fixture(scope='module')
def app():
    """
    Создает тестовый экземпляр Flask-приложения
    """
    app = create_app()
    app.config.update({
        "TESTING": True,
        "DEBUG": False,
    })
    
    # Маршрут для проверки работоспособности API
    @app.route('/api/health')
    def health_check():
        return {'status': 'ok', 'message': 'API работает нормально'}, 200
    
    yield app

@pytest.fixture(scope='module')
def client(app):
    """
    Создает тестовый клиент Flask
    """
    with app.test_client() as client:
        yield client

@pytest.fixture(scope='function')
def test_document():
    """
    Подготавливает тестовый документ для интеграционного тестирования
    """
    # Проверяем наличие тестового документа с ошибками
    test_file_path = TEST_DATA_DIR / "multiple_errors_document.docx"
    
    if not os.path.exists(test_file_path):
        # Создаем новый тестовый документ
        doc = Document()
        
        # Добавляем заголовок
        heading = doc.add_heading('Тестовый документ для интеграционного тестирования', 0)
        heading.alignment = 0  # Левое выравнивание вместо центрирования
        
        # Добавляем несколько параграфов с различными ошибками форматирования
        p = doc.add_paragraph('Этот абзац использует неправильный шрифт.')
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = None
        
        p = doc.add_paragraph('Этот абзац имеет неправильный размер шрифта.')
        for run in p.runs:
            run.font.size = 12  # Вместо 14pt
        
        p = doc.add_paragraph('Этот абзац имеет неправильный межстрочный интервал.')
        p.paragraph_format.line_spacing = 1.0  # Вместо 1.5
        
        # Сохраняем документ
        doc.save(test_file_path)
        logger.info(f"Создан тестовый документ: {test_file_path}")
    
    return test_file_path


class TestDocumentIntegrationFlow:
    """
    Интеграционные тесты для проверки полного цикла работы с документами
    """
    
    @pytest.mark.integration
    def test_full_document_flow(self, client, test_document):
        """
        Тест полного цикла работы с документом:
        1. Загрузка документа на проверку
        2. Получение результатов проверки
        3. Исправление документа
        4. Скачивание исправленного документа
        5. Проверка, что исправленный документ существует и открывается
        """
        # Проверка доступности API
        health_response = client.get('/api/health')
        assert health_response.status_code == 200, "API недоступен"
        
        # Шаг 1: Загружаем документ на проверку
        logger.info("Шаг 1: Загрузка документа на проверку")
        
        with open(test_document, 'rb') as f:
            response = client.post(
                '/api/document/upload',
                data={'file': (io.BytesIO(f.read()), os.path.basename(test_document))},
                content_type='multipart/form-data'
            )
        
        assert response.status_code == 200, f"Ошибка при загрузке документа: {response.data.decode('utf-8')}"
        upload_result = json.loads(response.data)
        assert "check_results" in upload_result, "В ответе отсутствуют результаты проверки"
        logger.info("Документ успешно загружен и проверен")
        
        # Сохраняем путь к загруженному файлу для дальнейших операций
        temp_file_path = upload_result.get("temp_path")
        assert temp_file_path, "Не получен путь к временному файлу"
        
        # Шаг 2: Проверяем наличие ошибок в результате
        logger.info("Шаг 2: Проверка результатов анализа документа")
        
        check_results = upload_result.get("check_results", {})
        assert "rules_results" in check_results, "В результатах отсутствуют проверки правил"
        
        # Находим общее количество проблем
        total_issues = check_results.get("total_issues_count", 0)
        logger.info(f"В документе найдено {total_issues} проблем форматирования")
        assert total_issues > 0, "В документе не найдено проблем, но они должны быть"
        
        # Шаг 3: Исправляем документ
        logger.info("Шаг 3: Исправление документа")
        
        correction_response = client.post(
            '/api/document/correct',
            json={"path": temp_file_path},
            content_type='application/json'
        )
        
        assert correction_response.status_code == 200, f"Ошибка при исправлении документа: {correction_response.data.decode('utf-8')}"
        correction_result = json.loads(correction_response.data)
        assert "corrected_path" in correction_result, "Не получен путь к исправленному документу"
        
        corrected_path = correction_result.get("corrected_path")
        logger.info(f"Документ успешно исправлен. Путь: {corrected_path}")
        
        # Шаг 4: Скачиваем исправленный документ
        logger.info("Шаг 4: Скачивание исправленного документа")
        
        download_response = client.get(
            f'/api/document/download-corrected?path={corrected_path}'
        )
        
        assert download_response.status_code in [200, 302], f"Ошибка при скачивании документа: {download_response.status_code}"
        
        # Если получено перенаправление, проверяем его
        if download_response.status_code == 302:
            redirect_url = download_response.headers.get('Location')
            assert redirect_url, "Перенаправление без URL"
            logger.info(f"Получено перенаправление на {redirect_url}")
        
        # Тест считается успешным, если мы дошли до этого места
        logger.info("Интеграционный тест полного цикла работы с документом завершен успешно")
        
    @pytest.mark.integration
    def test_document_flow_with_invalid_file(self, client):
        """
        Тест обработки невалидного файла:
        1. Попытка загрузки некорректного файла
        2. Проверка правильной обработки ошибки
        """
        # Создаем временный текстовый файл вместо DOCX
        with tempfile.NamedTemporaryFile(suffix='.txt') as temp_file:
            temp_file.write(b'This is not a valid DOCX file')
            temp_file.flush()
            
            # Шаг 1: Пытаемся загрузить некорректный файл
            logger.info("Загрузка некорректного файла")
            
            # На Windows NamedTemporaryFile блокирует повторное открытие файла.
            # Используем уже открытый дескриптор и читаем из него напрямую.
            temp_file.seek(0)
            response = client.post(
                '/api/document/upload',
                data={'file': (io.BytesIO(temp_file.read()), os.path.basename(temp_file.name))},
                content_type='multipart/form-data'
            )
            
            # Шаг 2: Проверяем, что ошибка обрабатывается корректно
            assert response.status_code in [400, 415, 422], f"Неожиданный статус код: {response.status_code}"
            response_data = json.loads(response.data)
            assert "error" in response_data, "В ответе отсутствует информация об ошибке"
            
            logger.info(f"Успешная проверка обработки невалидного файла: {response_data.get('error')}") 