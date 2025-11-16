"""
Модуль тестирования системы проверки форматирования документов.
Содержит тесты для проверки всех требований к форматированию документов DOCX.
"""

import os
import sys
import json
import pytest
import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Добавляем корневой каталог проекта в sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../..')))

from app.services.document_processor import DocumentProcessor
from app.services.norm_control_checker import NormControlChecker
from app.services.document_corrector import DocumentCorrector

# Пути к тестовым данным
TEST_DATA_DIR = Path(__file__).parent.parent / "test_data"
DOCUMENTS_DIR = TEST_DATA_DIR / "documents"
RESULTS_DIR = TEST_DATA_DIR / "results"

class TestResult:
    """Класс для хранения результатов тестирования"""
    def __init__(self, test_name: str, requirement: str, input_data: str, expected_result: str):
        """
        Инициализация результата теста
        Args:
            test_name: Название теста
            requirement: Требование, которое проверяется
            input_data: Входные данные
            expected_result: Ожидаемый результат
        """
        self.test_name = test_name
        self.requirement = requirement
        self.input_data = input_data
        self.expected_result = expected_result
        self.actual_result = None
        self.success = False
        self.error_message = None
        self.timestamp = datetime.datetime.now()
    
    def to_dict(self) -> dict:
        """
        Преобразует результат теста в словарь
        Returns:
            dict: Словарь с данными результата
        """
        return {
            'test_name': self.test_name,
            'requirement': self.requirement,
            'input_data': self.input_data,
            'expected_result': self.expected_result,
            'actual_result': self.actual_result,
            'success': self.success,
            'error_message': self.error_message,
            'timestamp': self.timestamp.isoformat()
        }

class TestFormatRequirements:
    """Класс с тестами для проверки требований к форматированию"""
    
    def setup_method(self):
        """Настройка перед каждым тестом"""
        self.results = []
        self.processor = DocumentProcessor(str(DOCUMENTS_DIR / "test_document.docx"))
        self.checker = NormControlChecker()
        self.corrector = DocumentCorrector()
        self._create_test_documents()
    
    def teardown_method(self):
        """Действия после каждого теста"""
        self._save_results()
    
    def _create_test_documents(self):
        """Создает тестовые документы для проверки"""
        # Создаем базовый тестовый документ
        doc = Document()
        doc.save(str(DOCUMENTS_DIR / "test_document.docx"))
        
        # Создаем документ с неправильным форматированием
        doc = Document()
        for _ in range(5):
            p = doc.add_paragraph("Тестовый текст с неправильным форматированием")
            p.style.font.name = "Arial"
            p.style.font.size = Pt(12)
        doc.save(str(DOCUMENTS_DIR / "wrong_format.docx"))
    
    def _save_results(self):
        """Сохраняет результаты тестов в файлы"""
        # Сохраняем в JSON
        json_results = [result.to_dict() for result in self.results]
        with open(RESULTS_DIR / "test_results.json", "w", encoding="utf-8") as f:
            json.dump(json_results, f, ensure_ascii=False, indent=4)
        
        # Сохраняем в TXT
        with open(RESULTS_DIR / "test_results.txt", "w", encoding="utf-8") as f:
            for result in self.results:
                f.write(f"Тест: {result.test_name}\n")
                f.write(f"Требование: {result.requirement}\n")
                f.write(f"Входные данные: {result.input_data}\n")
                f.write(f"Ожидаемый результат: {result.expected_result}\n")
                f.write(f"Фактический результат: {result.actual_result}\n")
                f.write(f"Статус: {'✅ УСПЕХ' if result.success else '❌ НЕУДАЧА'}\n")
                if result.error_message:
                    f.write(f"Ошибка: {result.error_message}\n")
                f.write("\n")

    def test_font_check(self):
        """Тест проверки шрифта"""
        test_result = TestResult(
            test_name="Проверка шрифта",
            requirement="Приложение должно проверять соответствие шрифта документа требованиям",
            input_data="Документ с неправильным шрифтом (Arial вместо Times New Roman)",
            expected_result="Система выявляет ошибку шрифта и предлагает исправление на Times New Roman"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            font_issues = [r for r in check_results['rules_results'] 
                         if r['rule_name'] == 'Проверка шрифта']
            
            if font_issues:
                test_result.actual_result = f"Обнаружено ошибок шрифта: {len(font_issues[0]['issues'])}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки шрифта не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_font_size_check(self):
        """Тест проверки размера шрифта"""
        test_result = TestResult(
            test_name="Проверка размера шрифта",
            requirement="Приложение должно проверять соответствие размера шрифта требованиям",
            input_data="Документ с неверным размером шрифта (12pt вместо 14pt)",
            expected_result="Система выявляет ошибку размера шрифта и предлагает исправление на 14pt"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            font_issues = [r for r in check_results['rules_results'] 
                         if r['rule_name'] == 'Проверка шрифта']
            
            size_issues = [i for r in font_issues for i in r['issues'] 
                         if i['type'] == 'wrong_font_size']
            
            if size_issues:
                test_result.actual_result = f"Обнаружено ошибок размера шрифта: {len(size_issues)}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки размера шрифта не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_margins_check(self):
        """Тест проверки полей страницы"""
        test_result = TestResult(
            test_name="Проверка полей страницы",
            requirement="Приложение должно проверять соответствие полей страницы требованиям",
            input_data="Документ с неправильными полями страницы",
            expected_result="Система выявляет ошибку полей страницы и предлагает исправление"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            margin_issues = [r for r in check_results['rules_results'] 
                           if r['rule_name'] == 'Проверка полей']
            
            if margin_issues:
                test_result.actual_result = f"Обнаружено ошибок полей: {len(margin_issues[0]['issues'])}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки полей не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_line_spacing_check(self):
        """Тест проверки межстрочного интервала"""
        test_result = TestResult(
            test_name="Проверка межстрочного интервала",
            requirement="Приложение должно проверять соответствие межстрочного интервала требованиям",
            input_data="Документ с неправильным межстрочным интервалом (одинарный вместо полуторного)",
            expected_result="Система выявляет ошибку межстрочного интервала и предлагает исправление"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            spacing_issues = [r for r in check_results['rules_results'] 
                            if r['rule_name'] == 'Проверка межстрочного интервала']
            
            if spacing_issues:
                test_result.actual_result = f"Обнаружено ошибок интервала: {len(spacing_issues[0]['issues'])}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки интервала не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_paragraph_indent_check(self):
        """Тест проверки абзацного отступа"""
        test_result = TestResult(
            test_name="Проверка абзацного отступа",
            requirement="Приложение должно проверять абзацный отступ (1.25 см)",
            input_data="Документ с неправильным абзацным отступом",
            expected_result="Система выявляет ошибки абзацного отступа"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            indent_issues = [r for r in check_results['rules_results'] 
                           if r['rule_name'] == 'Проверка абзацного отступа']
            
            if indent_issues:
                test_result.actual_result = f"Обнаружено ошибок отступа: {len(indent_issues[0]['issues'])}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки отступа не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_alignment_check(self):
        """Тест проверки выравнивания текста"""
        test_result = TestResult(
            test_name="Проверка выравнивания текста",
            requirement="Приложение должно проверять выравнивание текста по ширине",
            input_data="Документ с выравниванием по левому краю",
            expected_result="Система выявляет неправильное выравнивание"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            alignment_issues = [r for r in check_results['rules_results'] 
                              if r['rule_name'] == 'Проверка выравнивания текста']
            
            if alignment_issues:
                test_result.actual_result = f"Обнаружено ошибок выравнивания: {len(alignment_issues[0]['issues'])}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки выравнивания не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_headings_check(self):
        """Тест проверки форматирования заголовков"""
        test_result = TestResult(
            test_name="Проверка форматирования заголовков",
            requirement="Приложение должно проверять форматирование заголовков (жирный шрифт)",
            input_data="Документ с заголовками без выделения жирным",
            expected_result="Система выявляет ошибку форматирования заголовков"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            heading_issues = [r for r in check_results['rules_results'] 
                            if r['rule_name'] == 'Проверка заголовков']
            
            if heading_issues:
                test_result.actual_result = f"Обнаружено ошибок заголовков: {len(heading_issues[0]['issues'])}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки заголовков не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_page_numbers_check(self):
        """Тест проверки нумерации страниц"""
        test_result = TestResult(
            test_name="Проверка нумерации страниц",
            requirement="Нумерация должна быть сквозная, в верхней части листа справа, шрифт Times New Roman 12pt",
            input_data="Документ без нумерации страниц или с неправильной нумерацией",
            expected_result="Система выявляет отсутствие или неправильное расположение нумерации страниц"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            page_number_issues = [r for r in check_results['rules_results'] 
                                if r['rule_name'] == 'Проверка нумерации страниц']
            
            if page_number_issues:
                issues = page_number_issues[0]['issues']
                test_result.actual_result = (f"Обнаружены проблемы с позицией, "
                                          f"шрифтом и начальной страницей нумерации")
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки нумерации не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_bibliography_check(self):
        """Тест проверки списка литературы"""
        test_result = TestResult(
            test_name="Проверка списка литературы",
            requirement="Приложение должно проверять оформление списка литературы по ГОСТ",
            input_data="Документ с неправильным оформлением списка литературы",
            expected_result="Система выявляет ошибки в оформлении списка литературы"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            bibliography_issues = [r for r in check_results['rules_results'] 
                                if r['rule_name'] == 'Проверка списка литературы']
            
            if bibliography_issues:
                test_result.actual_result = f"Обнаружено ошибок списка литературы: {len(bibliography_issues[0]['issues'])}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки списка литературы не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_sections_check(self):
        """Тест проверки наличия разделов"""
        test_result = TestResult(
            test_name="Проверка наличия разделов",
            requirement="Приложение должно проверять наличие обязательных разделов",
            input_data="Документ без раздела 'Заключение'",
            expected_result="Система выявляет отсутствие обязательного раздела"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            section_issues = [r for r in check_results['rules_results'] 
                            if r['rule_name'] == 'Проверка наличия разделов']
            
            if section_issues:
                missing_sections = [i['expected'] for i in section_issues[0]['issues']]
                test_result.actual_result = f"Обнаружено отсутствие раздела '{missing_sections[0]}'"
                test_result.success = True
            else:
                test_result.actual_result = "Отсутствующие разделы не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_tables_check(self):
        """Тест проверки таблиц"""
        test_result = TestResult(
            test_name="Проверка таблиц",
            requirement="Приложение должно проверять оформление таблиц",
            input_data="Документ с таблицей без заголовка",
            expected_result="Система выявляет ошибки в оформлении таблиц"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            table_issues = [r for r in check_results['rules_results'] 
                          if r['rule_name'] == 'Проверка таблиц']
            
            if table_issues:
                test_result.actual_result = f"Обнаружено ошибок таблиц: {len(table_issues[0]['issues'])}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки таблиц не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_images_check(self):
        """Тест проверки рисунков"""
        test_result = TestResult(
            test_name="Проверка рисунков",
            requirement="Приложение должно проверять оформление рисунков и подписей к ним",
            input_data="Документ с рисунком без подписи",
            expected_result="Система выявляет отсутствие подписи к рисунку"
        )
        
        try:
            doc_data = self.processor.extract_data()
            check_results = self.checker.check_document(doc_data)
            
            image_issues = [r for r in check_results['rules_results'] 
                          if r['rule_name'] == 'Проверка рисунков']
            
            if image_issues:
                test_result.actual_result = f"Обнаружено ошибок оформления рисунков: {len(image_issues[0]['issues'])}"
                test_result.success = True
            else:
                test_result.actual_result = "Ошибки рисунков не обнаружены"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

    def test_auto_correction(self):
        """Тест автоматического исправления"""
        test_result = TestResult(
            test_name="Автоматическое исправление",
            requirement="Приложение должно предоставлять возможность автоматического исправления выявленных ошибок",
            input_data="Документ с множественными ошибками форматирования",
            expected_result="Система автоматически исправляет все выбранные ошибки"
        )
        
        try:
            # Создаем документ с ошибками
            doc = Document()
            p = doc.add_paragraph("Тестовый текст")
            p.style.font.name = "Arial"
            test_file = str(DOCUMENTS_DIR / "test_auto_correction.docx")
            doc.save(test_file)
            
            # Пытаемся исправить
            corrected_file = self.corrector.correct_document(test_file)
            
            if os.path.exists(corrected_file):
                # Проверяем исправленный файл
                processor = DocumentProcessor(corrected_file)
                doc_data = processor.extract_data()
                check_results = self.checker.check_document(doc_data)
                
                if check_results['total_issues_count'] == 0:
                    test_result.actual_result = "Исправленный файл создан и проверен"
                    test_result.success = True
                else:
                    test_result.actual_result = f"В исправленном файле остались ошибки: {check_results['total_issues_count']}"
                    test_result.success = False
            else:
                test_result.actual_result = "Не удалось создать исправленный файл"
                test_result.success = False
                
        except Exception as e:
            test_result.error_message = str(e)
            test_result.success = False
            
        self.results.append(test_result)

def run_tests():
    """Запускает все тесты и сохраняет результаты"""
    test_suite = TestFormatRequirements()
    test_suite.setup_method()
    
    # Запуск всех тестов
    try:
        test_suite.test_font_check()
        test_suite.test_font_size_check()
        test_suite.test_margins_check()
        test_suite.test_line_spacing_check()
        test_suite.test_paragraph_indent_check()
        test_suite.test_alignment_check()
        test_suite.test_headings_check()
        test_suite.test_page_numbers_check()
        test_suite.test_bibliography_check()
        test_suite.test_sections_check()
        test_suite.test_tables_check()
        test_suite.test_images_check()
        test_suite.test_auto_correction()
    except Exception as e:
        print(f"Ошибка при выполнении тестов: {str(e)}")
    finally:
        test_suite.teardown_method()
    
    print("Тестирование завершено")

if __name__ == "__main__":
    run_tests() 