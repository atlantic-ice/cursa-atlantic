"""
Модуль с основными классами и методами для проверки форматирования документов.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
from typing import Dict, List, Any, Optional

class DocumentProcessor:
    """Класс для извлечения данных из документа"""
    
    def __init__(self, file_path: str):
        """
        Инициализация процессора документов
        Args:
            file_path (str): Путь к документу DOCX
        """
        self.file_path = file_path
        self.document = Document(file_path)
        
    def extract_data(self) -> Dict[str, Any]:
        """
        Извлекает все необходимые данные из документа
        Returns:
            Dict[str, Any]: Словарь с данными документа
        """
        return {
            'paragraphs': self._extract_paragraphs(),
            'tables': self._extract_tables(),
            'headings': self._extract_headings(),
            'bibliography': self._extract_bibliography(),
            'styles': self._extract_styles(),
            'page_setup': self._extract_page_setup(),
            'images': self._extract_images(),
            'page_numbers': self._extract_page_numbers(),
            'sections': self._extract_sections(),
            'lists': self._extract_lists(),
            'formulas': self._extract_formulas(),
            'cross_references': self._extract_cross_references(),
            'appendices': self._extract_appendices()
        }

    def _extract_paragraphs(self) -> List[Dict[str, Any]]:
        """
        Извлекает информацию о параграфах
        Returns:
            List[Dict[str, Any]]: Список параграфов с их свойствами
        """
        paragraphs_data = []
        for paragraph in self.document.paragraphs:
            if not paragraph.text.strip():
                continue
                
            paragraph_data = {
                'text': paragraph.text,
                'style_name': paragraph.style.name if paragraph.style else None,
                'alignment': paragraph.alignment,
                'line_spacing': paragraph.paragraph_format.line_spacing,
                'line_spacing_rule': paragraph.paragraph_format.line_spacing_rule,
                'first_line_indent': paragraph.paragraph_format.first_line_indent.cm if paragraph.paragraph_format.first_line_indent else None,
                'font_data': self._extract_font_data(paragraph),
                'is_heading': paragraph.style.name.startswith('Heading') if paragraph.style else False
            }
            paragraphs_data.append(paragraph_data)
        return paragraphs_data

    def _extract_font_data(self, paragraph) -> Dict[str, Any]:
        """
        Извлекает информацию о шрифте из параграфа
        Args:
            paragraph: Параграф документа
        Returns:
            Dict[str, Any]: Информация о шрифте
        """
        runs = paragraph.runs
        if not runs:
            return {}
            
        # Берем первый run как основной для параграфа
        main_run = runs[0]
        font_data = {
            'name': main_run.font.name,
            'size': main_run.font.size.pt if main_run.font.size else None,
            'bold': main_run.font.bold,
            'italic': main_run.font.italic,
            'underline': main_run.font.underline,
            'color': self._get_font_color(main_run),
            'all_caps': main_run.font.all_caps,
            'variations': self._get_font_variations(runs)
        }
        return font_data

    def _get_font_color(self, run) -> Optional[str]:
        """
        Получает цвет шрифта в формате RGB
        Args:
            run: Элемент текста
        Returns:
            Optional[str]: Цвет в формате RGB или None
        """
        if not run.font.color.rgb:
            return None
        return run.font.color.rgb

    def _get_font_variations(self, runs) -> List[Dict[str, Any]]:
        """
        Получает все вариации форматирования текста в параграфе
        Args:
            runs: Список элементов текста
        Returns:
            List[Dict[str, Any]]: Список вариаций форматирования
        """
        variations = []
        for run in runs[1:]:  # Пропускаем первый run, так как он основной
            if run.text.strip():
                variation = {
                    'text': run.text,
                    'font_name': run.font.name,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline,
                    'color': self._get_font_color(run)
                }
                variations.append(variation)
        return variations

    def _extract_tables(self) -> List[Dict[str, Any]]:
        """
        Извлекает информацию о таблицах
        Returns:
            List[Dict[str, Any]]: Список таблиц с их свойствами
        """
        tables_data = []
        for table in self.document.tables:
            table_data = {
                'row_count': len(table.rows),
                'column_count': len(table.columns),
                'style': table.style.name if table.style else None,
                'cells': self._extract_table_cells(table),
                'has_header': self._check_table_header(table),
                'caption': self._find_table_caption(table)
            }
            tables_data.append(table_data)
        return tables_data

    def _extract_table_cells(self, table) -> List[List[Dict[str, Any]]]:
        """
        Извлекает данные ячеек таблицы
        Args:
            table: Таблица документа
        Returns:
            List[List[Dict[str, Any]]]: Матрица данных ячеек
        """
        cells_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_data = {
                    'text': cell.text,
                    'paragraphs': [self._extract_font_data(p) for p in cell.paragraphs if p.text.strip()],
                    'width': cell.width.cm if cell.width else None,
                    'vertical_alignment': cell._tc.get_or_add_tcPr().get_or_add_vAlign().val
                }
                row_data.append(cell_data)
            cells_data.append(row_data)
        return cells_data

    def _check_table_header(self, table) -> bool:
        """
        Проверяет наличие заголовка таблицы
        Args:
            table: Таблица документа
        Returns:
            bool: True если есть заголовок
        """
        if not table.rows:
            return False
        
        first_row = table.rows[0]
        # Проверяем форматирование первой строки
        return any(cell._tc.get_or_add_tcPr().get_or_add_vAlign().val == 'center' 
                  or self._has_bold_text(cell) for cell in first_row.cells)

    def _has_bold_text(self, cell) -> bool:
        """
        Проверяет наличие жирного текста в ячейке
        Args:
            cell: Ячейка таблицы
        Returns:
            bool: True если есть жирный текст
        """
        return any(run.font.bold for paragraph in cell.paragraphs 
                  for run in paragraph.runs)

    def _find_table_caption(self, table) -> Optional[str]:
        """
        Ищет подпись к таблице
        Args:
            table: Таблица документа
        Returns:
            Optional[str]: Текст подписи или None
        """
        # Ищем параграф перед таблицей
        for paragraph in self.document.paragraphs:
            if paragraph._p.getnext() is not None and table._tbl in paragraph._p.getnext():
                if paragraph.text.strip().lower().startswith('таблица'):
                    return paragraph.text.strip()
        return None

    def _extract_headings(self) -> List[Dict[str, Any]]:
        """
        Извлекает информацию о заголовках
        Returns:
            List[Dict[str, Any]]: Список заголовков с их свойствами
        """
        headings_data = []
        for paragraph in self.document.paragraphs:
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                heading_data = {
                    'level': int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 0,
                    'text': paragraph.text,
                    'font_data': self._extract_font_data(paragraph),
                    'alignment': paragraph.alignment,
                    'page_number': self._get_paragraph_page(paragraph)
                }
                headings_data.append(heading_data)
        return headings_data

    def _get_paragraph_page(self, paragraph) -> int:
        """
        Определяет номер страницы параграфа
        Args:
            paragraph: Параграф документа
        Returns:
            int: Номер страницы
        """
        section = paragraph._p.getparent().getparent()
        page_count = 1
        for sect in self.document.sections:
            if sect._sectPr == section:
                break
            page_count += 1
        return page_count

    def _extract_bibliography(self) -> Dict[str, Any]:
        """
        Извлекает информацию о списке литературы
        Returns:
            Dict[str, Any]: Информация о списке литературы
        """
        bibliography_data = {
            'exists': False,
            'items': [],
            'format': None,
            'location': None
        }
        
        # Ищем раздел со списком литературы
        for i, paragraph in enumerate(self.document.paragraphs):
            if any(title in paragraph.text.lower() for title in ['список литературы', 'библиографический список']):
                bibliography_data['exists'] = True
                bibliography_data['location'] = i
                
                # Собираем элементы списка литературы
                items = []
                current_item = ''
                for p in self.document.paragraphs[i+1:]:
                    if p.style and p.style.name.startswith('Heading'):
                        break
                    if p.text.strip():
                        if re.match(r'^\d+\.', p.text):
                            if current_item:
                                items.append(current_item)
                            current_item = p.text
                        else:
                            current_item += ' ' + p.text
                
                if current_item:
                    items.append(current_item)
                
                bibliography_data['items'] = items
                bibliography_data['format'] = self._detect_bibliography_format(items)
                break
        
        return bibliography_data

    def _detect_bibliography_format(self, items: List[str]) -> str:
        """
        Определяет формат оформления списка литературы
        Args:
            items: Список элементов библиографии
        Returns:
            str: Формат оформления (ГОСТ/другой)
        """
        gost_patterns = [
            r'\d+\.\s+[А-Я][а-я]+\s+[А-Я]\.[А-Я]\.',  # Автор И.О.
            r'\d+\.\s+[А-Я][а-я]+,\s+[А-Я]\.[А-Я]\.',  # Автор, И.О.
            r'[МСП][\.:]',  # Город издания
            r'\d{4}\.',  # Год издания
            r'\d+\s+с\.'  # Количество страниц
        ]
        
        gost_matches = 0
        for item in items:
            matches = sum(1 for pattern in gost_patterns if re.search(pattern, item))
            if matches >= 3:  # Если найдено более 3 признаков ГОСТа
                gost_matches += 1
        
        return 'ГОСТ' if gost_matches >= len(items) * 0.7 else 'Другой'

    def _extract_styles(self) -> Dict[str, Any]:
        """
        Извлекает информацию о стилях документа
        Returns:
            Dict[str, Any]: Информация о стилях
        """
        styles_data = {
            'paragraph_styles': [],
            'character_styles': [],
            'table_styles': []
        }
        
        for style in self.document.styles:
            style_data = {
                'name': style.name,
                'type': style.type,
                'font': {
                    'name': style.font.name if hasattr(style, 'font') else None,
                    'size': style.font.size.pt if hasattr(style, 'font') and style.font.size else None,
                    'bold': style.font.bold if hasattr(style, 'font') else None,
                    'italic': style.font.italic if hasattr(style, 'font') else None
                } if hasattr(style, 'font') else None
            }
            
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                styles_data['paragraph_styles'].append(style_data)
            elif style.type == WD_STYLE_TYPE.CHARACTER:
                styles_data['character_styles'].append(style_data)
            elif style.type == WD_STYLE_TYPE.TABLE:
                styles_data['table_styles'].append(style_data)
        
        return styles_data

    def _extract_page_setup(self) -> Dict[str, float]:
        """
        Извлекает настройки страницы
        Returns:
            Dict[str, float]: Настройки страницы в сантиметрах
        """
        section = self.document.sections[0]
        return {
            'top_margin': section.top_margin.cm,
            'bottom_margin': section.bottom_margin.cm,
            'left_margin': section.left_margin.cm,
            'right_margin': section.right_margin.cm,
            'page_height': section.page_height.cm,
            'page_width': section.page_width.cm,
            'header_distance': section.header_distance.cm,
            'footer_distance': section.footer_distance.cm,
            'orientation': 'portrait' if section.orientation == 0 else 'landscape'
        }

    def _extract_images(self) -> List[Dict[str, Any]]:
        """
        Извлекает информацию о изображениях
        Returns:
            List[Dict[str, Any]]: Список изображений с их свойствами
        """
        images_data = []
        for rel in self.document.part.rels.values():
            if "image" in rel.reltype:
                # Ищем параграф с подписью к рисунку
                caption = None
                for paragraph in self.document.paragraphs:
                    if rel.target_ref in paragraph._element.xml:
                        # Ищем подпись в следующем параграфе
                        next_p = paragraph._p.getnext()
                        if next_p is not None:
                            next_text = next_p.text if hasattr(next_p, 'text') else ''
                            if next_text.strip().lower().startswith('рис'):
                                caption = next_text.strip()
                
                image_data = {
                    'filename': rel.target_ref,
                    'type': rel.reltype,
                    'caption': caption,
                    'has_caption': caption is not None
                }
                images_data.append(image_data)
        return images_data

    def _extract_page_numbers(self) -> Dict[str, Any]:
        """
        Извлекает информацию о нумерации страниц
        Returns:
            Dict[str, Any]: Информация о нумерации страниц
        """
        page_numbers = {
            'has_page_numbers': False,
            'position': None,
            'first_numbered_page': None,
            'alignment': None,
            'font': {
                'name': None,
                'size': None,
                'bold': None
            }
        }
        
        try:
            # Проверяем колонтитулы на наличие нумерации
            for section in self.document.sections:
                # Проверка верхнего колонтитула
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph._element.xpath('.//w:fldChar'):
                            field_codes = paragraph._element.xpath('.//w:instrText')
                            for field in field_codes:
                                if 'PAGE' in field.text:
                                    page_numbers.update({
                                        'has_page_numbers': True,
                                        'position': 'header',
                                        'alignment': paragraph.alignment,
                                        'font': self._extract_font_data(paragraph)
                                    })
                                    return page_numbers
                
                # Проверка нижнего колонтитула
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph._element.xpath('.//w:fldChar'):
                            field_codes = paragraph._element.xpath('.//w:instrText')
                            for field in field_codes:
                                if 'PAGE' in field.text:
                                    page_numbers.update({
                                        'has_page_numbers': True,
                                        'position': 'footer',
                                        'alignment': paragraph.alignment,
                                        'font': self._extract_font_data(paragraph)
                                    })
                                    return page_numbers
        except Exception as e:
            print(f"Ошибка при проверке нумерации страниц: {str(e)}")
        
        return page_numbers

class NormControlChecker:
    """Класс для проверки документа на соответствие требованиям нормоконтроля"""
    
    def __init__(self):
        """Инициализация проверяющего с правилами оформления"""
        self.standard_rules = {
            'font': {
                'name': 'Times New Roman',
                'size': 14.0
            },
            'margins': {
                'left': 3.0,
                'right': 1.5,
                'top': 2.0,
                'bottom': 2.0
            },
            'line_spacing': 1.5,
            'paragraph_indent': 1.25,
            'alignment': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            'heading_alignment': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'page_numbers': {
                'position': 'header',
                'alignment': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'font_name': 'Times New Roman',
                'font_size': 12.0
            },
            'required_sections': [
                'введение',
                'заключение',
                'список литературы'
            ],
            'bibliography_format': 'ГОСТ'
        }

    def check_document(self, document_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Проверяет документ на соответствие требованиям
        Args:
            document_data: Данные документа от DocumentProcessor
        Returns:
            Dict[str, Any]: Результаты проверки
        """
        results = []
        
        # Проверка шрифта
        font_issues = self._check_font(document_data)
        if font_issues:
            results.append({
                'rule_name': 'Проверка шрифта',
                'issues': font_issues
            })
        
        # Проверка полей
        margin_issues = self._check_margins(document_data)
        if margin_issues:
            results.append({
                'rule_name': 'Проверка полей',
                'issues': margin_issues
            })
        
        # Проверка межстрочного интервала
        spacing_issues = self._check_line_spacing(document_data)
        if spacing_issues:
            results.append({
                'rule_name': 'Проверка межстрочного интервала',
                'issues': spacing_issues
            })
        
        # Проверка абзацного отступа
        indent_issues = self._check_paragraph_indent(document_data)
        if indent_issues:
            results.append({
                'rule_name': 'Проверка абзацного отступа',
                'issues': indent_issues
            })
        
        # Проверка выравнивания
        alignment_issues = self._check_alignment(document_data)
        if alignment_issues:
            results.append({
                'rule_name': 'Проверка выравнивания текста',
                'issues': alignment_issues
            })
        
        # Проверка заголовков
        heading_issues = self._check_headings(document_data)
        if heading_issues:
            results.append({
                'rule_name': 'Проверка заголовков',
                'issues': heading_issues
            })
        
        # Проверка нумерации страниц
        page_number_issues = self._check_page_numbers(document_data)
        if page_number_issues:
            results.append({
                'rule_name': 'Проверка нумерации страниц',
                'issues': page_number_issues
            })
        
        # Проверка списка литературы
        bibliography_issues = self._check_bibliography(document_data)
        if bibliography_issues:
            results.append({
                'rule_name': 'Проверка списка литературы',
                'issues': bibliography_issues
            })
        
        # Проверка наличия разделов
        section_issues = self._check_required_sections(document_data)
        if section_issues:
            results.append({
                'rule_name': 'Проверка наличия разделов',
                'issues': section_issues
            })
        
        # Проверка таблиц
        table_issues = self._check_tables(document_data)
        if table_issues:
            results.append({
                'rule_name': 'Проверка таблиц',
                'issues': table_issues
            })
        
        # Проверка рисунков
        image_issues = self._check_images(document_data)
        if image_issues:
            results.append({
                'rule_name': 'Проверка рисунков',
                'issues': image_issues
            })
        
        return {
            'rules_results': results,
            'total_issues_count': sum(len(r['issues']) for r in results)
        }

    def _check_font(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет соответствие шрифта требованиям
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с шрифтом
        """
        issues = []
        for paragraph in document_data['paragraphs']:
            font_data = paragraph['font_data']
            if font_data.get('name') != self.standard_rules['font']['name']:
                issues.append({
                    'type': 'wrong_font',
                    'location': f"Текст: {paragraph['text'][:50]}...",
                    'found': font_data.get('name'),
                    'expected': self.standard_rules['font']['name'],
                    'severity': 'high',
                    'auto_fixable': True
                })
            
            if font_data.get('size') != self.standard_rules['font']['size']:
                issues.append({
                    'type': 'wrong_font_size',
                    'location': f"Текст: {paragraph['text'][:50]}...",
                    'found': font_data.get('size'),
                    'expected': self.standard_rules['font']['size'],
                    'severity': 'high',
                    'auto_fixable': True
                })
        return issues

    def _check_margins(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет соответствие полей страницы требованиям
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с полями
        """
        issues = []
        page_setup = document_data['page_setup']
        margins = self.standard_rules['margins']
        
        for side in ['left', 'right', 'top', 'bottom']:
            margin_key = f'{side}_margin'
            if abs(page_setup[margin_key] - margins[side]) > 0.1:  # Допуск 1мм
                issues.append({
                    'type': f'wrong_{side}_margin',
                    'location': 'Документ',
                    'found': page_setup[margin_key],
                    'expected': margins[side],
                    'severity': 'medium',
                    'auto_fixable': True
                })
        return issues

    def _check_line_spacing(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет межстрочный интервал
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с интервалами
        """
        issues = []
        for paragraph in document_data['paragraphs']:
            if paragraph.get('line_spacing') != self.standard_rules['line_spacing']:
                issues.append({
                    'type': 'wrong_line_spacing',
                    'location': f"Текст: {paragraph['text'][:50]}...",
                    'found': paragraph.get('line_spacing'),
                    'expected': self.standard_rules['line_spacing'],
                    'severity': 'medium',
                    'auto_fixable': True
                })
        return issues

    def _check_paragraph_indent(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет абзацный отступ
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с отступами
        """
        issues = []
        for paragraph in document_data['paragraphs']:
            if not paragraph.get('is_heading'):  # Пропускаем заголовки
                indent = paragraph.get('first_line_indent')
                if abs(indent - self.standard_rules['paragraph_indent']) > 0.1:  # Допуск 1мм
                    issues.append({
                        'type': 'wrong_paragraph_indent',
                        'location': f"Текст: {paragraph['text'][:50]}...",
                        'found': indent,
                        'expected': self.standard_rules['paragraph_indent'],
                        'severity': 'low',
                        'auto_fixable': True
                    })
        return issues

    def _check_alignment(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет выравнивание текста
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с выравниванием
        """
        issues = []
        for paragraph in document_data['paragraphs']:
            if paragraph.get('alignment') != self.standard_rules['alignment']:
                issues.append({
                    'type': 'wrong_alignment',
                    'location': f"Текст: {paragraph['text'][:50]}...",
                    'found': paragraph.get('alignment'),
                    'expected': self.standard_rules['alignment'],
                    'severity': 'medium',
                    'auto_fixable': True
                })
        return issues

    def _check_headings(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет соответствие заголовков требованиям
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с заголовками
        """
        issues = []
        for heading in document_data['headings']:
            if heading.get('alignment') != self.standard_rules['heading_alignment']:
                issues.append({
                    'type': 'wrong_heading_alignment',
                    'location': f"Заголовок: {heading['text'][:50]}...",
                    'found': heading.get('alignment'),
                    'expected': self.standard_rules['heading_alignment'],
                    'severity': 'medium',
                    'auto_fixable': True
                })
        return issues

    def _check_page_numbers(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет соответствие нумерации страниц требованиям
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с нумерацией страниц
        """
        issues = []
        page_numbers = document_data['page_numbers']
        if page_numbers.get('has_page_numbers') != True:
            issues.append({
                'type': 'missing_page_numbers',
                'location': 'Документ',
                'severity': 'high',
                'auto_fixable': False
            })
        elif page_numbers.get('position') != self.standard_rules['page_numbers']['position']:
            issues.append({
                'type': 'wrong_page_numbers_position',
                'location': 'Документ',
                'found': page_numbers['position'],
                'expected': self.standard_rules['page_numbers']['position'],
                'severity': 'medium',
                'auto_fixable': True
            })
        elif page_numbers.get('alignment') != self.standard_rules['page_numbers']['alignment']:
            issues.append({
                'type': 'wrong_page_numbers_alignment',
                'location': 'Документ',
                'found': page_numbers['alignment'],
                'expected': self.standard_rules['page_numbers']['alignment'],
                'severity': 'medium',
                'auto_fixable': True
            })
        elif page_numbers.get('font_name') != self.standard_rules['page_numbers']['font_name']:
            issues.append({
                'type': 'wrong_page_numbers_font_name',
                'location': 'Документ',
                'found': page_numbers['font_name'],
                'expected': self.standard_rules['page_numbers']['font_name'],
                'severity': 'medium',
                'auto_fixable': True
            })
        elif page_numbers.get('font_size') != self.standard_rules['page_numbers']['font_size']:
            issues.append({
                'type': 'wrong_page_numbers_font_size',
                'location': 'Документ',
                'found': page_numbers['font_size'],
                'expected': self.standard_rules['page_numbers']['font_size'],
                'severity': 'medium',
                'auto_fixable': True
            })
        return issues

    def _check_bibliography(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет соответствие списка литературы требованиям
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с списком литературы
        """
        issues = []
        bibliography = document_data['bibliography']
        if bibliography.get('exists') != True:
            issues.append({
                'type': 'missing_bibliography',
                'location': 'Документ',
                'severity': 'high',
                'auto_fixable': False
            })
        elif bibliography.get('format') != self.standard_rules['bibliography_format']:
            issues.append({
                'type': 'wrong_bibliography_format',
                'location': 'Документ',
                'found': bibliography['format'],
                'expected': self.standard_rules['bibliography_format'],
                'severity': 'high',
                'auto_fixable': False
            })
        return issues

    def _check_required_sections(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет наличие необходимых разделов в документе
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с разделами
        """
        issues = []
        for section in self.standard_rules['required_sections']:
            if section not in document_data['sections']:
                issues.append({
                    'type': 'missing_section',
                    'location': f"Документ",
                    'expected': section,
                    'severity': 'high',
                    'auto_fixable': False
                })
        return issues

    def _check_tables(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет соответствие таблиц требованиям
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с таблицами
        """
        issues = []
        for table in document_data['tables']:
            if table.get('style') != None and table.get('style') not in self.standard_rules['table_styles']:
                issues.append({
                    'type': 'wrong_table_style',
                    'location': f"Таблица: {table['caption'][:50]}...",
                    'found': table['style'],
                    'severity': 'medium',
                    'auto_fixable': False
                })
        return issues

    def _check_images(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Проверяет соответствие рисунков требованиям
        Args:
            document_data: Данные документа
        Returns:
            List[Dict[str, Any]]: Список проблем с рисунками
        """
        issues = []
        for image in document_data['images']:
            if image.get('type') != 'image':
                issues.append({
                    'type': 'wrong_image_type',
                    'location': f"Изображение: {image['filename']}",
                    'found': image['type'],
                    'severity': 'medium',
                    'auto_fixable': False
                })
        return issues

class DocumentCorrector:
    """Класс для исправления ошибок форматирования в документе"""
    
    def __init__(self):
        """Инициализация корректора документов"""
        self.standard_rules = {
            'font': {
                'name': 'Times New Roman',
                'size': 14.0
            },
            'margins': {
                'left': 3.0,
                'right': 1.5,
                'top': 2.0,
                'bottom': 2.0
            },
            'line_spacing': 1.5,
            'paragraph_indent': 1.25,
            'alignment': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            'heading_alignment': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'page_numbers': {
                'position': 'header',
                'alignment': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'font_name': 'Times New Roman',
                'font_size': 12.0
            }
        }

    def correct_document(self, file_path: str) -> str:
        """
        Исправляет ошибки форматирования в документе
        Args:
            file_path: Путь к документу
        Returns:
            str: Путь к исправленному документу
        """
        doc = Document(file_path)
        
        # Исправление шрифта
        self._correct_font(doc)
        
        # Исправление полей
        self._correct_margins(doc)
        
        # Исправление межстрочного интервала
        self._correct_line_spacing(doc)
        
        # Исправление абзацного отступа
        self._correct_paragraph_indent(doc)
        
        # Исправление выравнивания
        self._correct_alignment(doc)
        
        # Исправление заголовков
        self._correct_headings(doc)
        
        # Исправление нумерации страниц
        self._correct_page_numbers(doc)
        
        # Сохранение исправленного документа
        output_path = self._get_output_path(file_path)
        doc.save(output_path)
        return output_path

    def _correct_font(self, document: Document) -> None:
        """
        Исправляет шрифт во всем документе
        Args:
            document: Документ Word
        """
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.name = self.standard_rules['font']['name']
                run.font.size = Pt(self.standard_rules['font']['size'])

    def _correct_margins(self, document: Document) -> None:
        """
        Исправляет поля страницы
        Args:
            document: Документ Word
        """
        for section in document.sections:
            section.left_margin = Cm(self.standard_rules['margins']['left'])
            section.right_margin = Cm(self.standard_rules['margins']['right'])
            section.top_margin = Cm(self.standard_rules['margins']['top'])
            section.bottom_margin = Cm(self.standard_rules['margins']['bottom'])

    def _correct_line_spacing(self, document: Document) -> None:
        """
        Исправляет межстрочный интервал
        Args:
            document: Документ Word
        """
        for paragraph in document.paragraphs:
            paragraph.paragraph_format.line_spacing = self.standard_rules['line_spacing']
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    def _correct_paragraph_indent(self, document: Document) -> None:
        """
        Исправляет абзацный отступ
        Args:
            document: Документ Word
        """
        for paragraph in document.paragraphs:
            if not paragraph.style or not paragraph.style.name.startswith('Heading'):
                paragraph.paragraph_format.first_line_indent = Cm(self.standard_rules['paragraph_indent'])

    def _correct_alignment(self, document: Document) -> None:
        """
        Исправляет выравнивание текста
        Args:
            document: Документ Word
        """
        for paragraph in document.paragraphs:
            if not paragraph.style or not paragraph.style.name.startswith('Heading'):
                paragraph.alignment = self.standard_rules['alignment']

    def _correct_headings(self, document: Document) -> None:
        """
        Исправляет форматирование заголовков
        Args:
            document: Документ Word
        """
        for paragraph in document.paragraphs:
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                paragraph.alignment = self.standard_rules['heading_alignment']
                for run in paragraph.runs:
                    run.font.bold = True

    def _correct_page_numbers(self, document: Document) -> None:
        """
        Исправляет нумерацию страниц
        Args:
            document: Документ Word
        """
        # Удаляем существующую нумерацию
        for section in document.sections:
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph._element.xpath('.//w:fldChar'):
                        p = paragraph._element
                        p.getparent().remove(p)
            
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph._element.xpath('.//w:fldChar'):
                        p = paragraph._element
                        p.getparent().remove(p)
        
        # Добавляем новую нумерацию в верхний колонтитул
        for section in document.sections:
            # Отключаем связь с предыдущим колонтитулом
            section.header.is_linked_to_previous = False
            
            # Создаем параграф для номера страницы
            header_paragraph = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
            header_paragraph.alignment = self.standard_rules['page_numbers']['alignment']
            
            # Добавляем номер страницы
            run = header_paragraph.add_run()
            run.font.name = self.standard_rules['page_numbers']['font_name']
            run.font.size = Pt(self.standard_rules['page_numbers']['font_size'])
            self._add_page_number(run)

    def _add_page_number(self, run) -> None:
        """
        Добавляет номер страницы через поле
        Args:
            run: Элемент текста для добавления номера
        """
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE \\* MERGEFORMAT "

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "1"  # Placeholder для номера страницы

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(fldChar3)
        run._element.append(fldChar4)

    def _get_output_path(self, input_path: str) -> str:
        """
        Генерирует путь для сохранения исправленного документа
        Args:
            input_path: Путь к исходному документу
        Returns:
            str: Путь к исправленному документу
        """
        dir_name = os.path.dirname(input_path)
        file_name = os.path.basename(input_path)
        name, ext = os.path.splitext(file_name)
        return os.path.join(dir_name, f"{name}_corrected{ext}") 