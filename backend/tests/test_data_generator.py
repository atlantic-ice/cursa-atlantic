#!/usr/bin/env python
"""
Модуль для подготовки тестовых DOCX-файлов
"""
import os
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# Добавляем путь к корневой директории проекта
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Директории для тестовых файлов
TEST_DATA_DIR = Path(__file__).parent / "test_data"
os.makedirs(TEST_DATA_DIR, exist_ok=True)

class TestDocumentGenerator:
    """
    Класс для генерации тестовых документов DOCX
    """
    
    def __init__(self):
        """
        Инициализация генератора тестовых данных
        """
        self.document_files = {}  # Словарь для хранения путей к сгенерированным файлам
    
    def create_empty_document(self):
        """
        Создает пустой тестовый документ
        
        Returns:
            str: Путь к созданному файлу
        """
        file_path = TEST_DATA_DIR / "empty_document.docx"
        
        doc = Document()
        doc.add_paragraph("Тестовый документ")
        doc.save(file_path)
        
        self.document_files['empty'] = file_path
        print(f"Создан пустой документ: {file_path}")
        return file_path
    
    def create_wrong_font_document(self):
        """
        Создает документ с неправильным шрифтом
        
        Returns:
            str: Путь к созданному файлу
        """
        file_path = TEST_DATA_DIR / "wrong_font_document.docx"
        
        doc = Document()
        paragraph = doc.add_paragraph("Текст с неправильным шрифтом (Arial вместо Times New Roman)")
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(14)
        
        doc.save(file_path)
        
        self.document_files['wrong_font'] = file_path
        print(f"Создан документ с неправильным шрифтом: {file_path}")
        return file_path
    
    def create_wrong_font_size_document(self):
        """
        Создает документ с неправильным размером шрифта
        
        Returns:
            str: Путь к созданному файлу
        """
        file_path = TEST_DATA_DIR / "wrong_font_size_document.docx"
        
        doc = Document()
        paragraph = doc.add_paragraph("Текст с неправильным размером шрифта (12pt вместо 14pt)")
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        doc.save(file_path)
        
        self.document_files['wrong_font_size'] = file_path
        print(f"Создан документ с неправильным размером шрифта: {file_path}")
        return file_path
    
    def create_wrong_margins_document(self):
        """
        Создает документ с неправильными полями страницы
        
        Returns:
            str: Путь к созданному файлу
        """
        file_path = TEST_DATA_DIR / "wrong_margins_document.docx"
        
        doc = Document()
        # Настройка полей (не соответствует требованиям)
        section = doc.sections[0]
        section.left_margin = Cm(2)  # Вместо 3 см
        section.right_margin = Cm(2)  # Вместо 1.5 см
        section.top_margin = Cm(1.5)  # Вместо 2 см
        section.bottom_margin = Cm(1.5)  # Вместо 2 см
        
        doc.add_paragraph("Документ с неправильными полями страницы")
        doc.save(file_path)
        
        self.document_files['wrong_margins'] = file_path
        print(f"Создан документ с неправильными полями: {file_path}")
        return file_path
    
    def create_wrong_line_spacing_document(self):
        """
        Создает документ с неправильным межстрочным интервалом
        
        Returns:
            str: Путь к созданному файлу
        """
        file_path = TEST_DATA_DIR / "wrong_line_spacing_document.docx"
        
        doc = Document()
        paragraph = doc.add_paragraph("Текст с одинарным межстрочным интервалом (1.0 вместо 1.5)")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.0  # Одинарный интервал вместо 1.5
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        doc.save(file_path)
        
        self.document_files['wrong_line_spacing'] = file_path
        print(f"Создан документ с неправильным межстрочным интервалом: {file_path}")
        return file_path
    
    def create_multiple_errors_document(self):
        """
        Создает документ с множественными ошибками форматирования
        
        Returns:
            str: Путь к созданному файлу
        """
        file_path = TEST_DATA_DIR / "multiple_errors_document.docx"
        
        doc = Document()
        
        # Неверные поля
        section = doc.sections[0]
        section.left_margin = Cm(2)  # Вместо 3 см
        section.right_margin = Cm(2)  # Вместо 1.5 см
        section.top_margin = Cm(1.5)  # Вместо 2 см
        section.bottom_margin = Cm(1.5)  # Вместо 2 см
        
        # Заголовок с неверным форматированием
        heading = doc.add_heading('Документ с множественными ошибками форматирования', level=1)
        for run in heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
        
        # Основной текст с неверным шрифтом и размером
        paragraph1 = doc.add_paragraph("Этот абзац имеет неправильный шрифт (Arial) и размер (12pt).")
        for run in paragraph1.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
        
        # Неверный межстрочный интервал
        paragraph1.paragraph_format.line_spacing = 1.0
        paragraph1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Неверное выравнивание
        paragraph2 = doc.add_paragraph("Этот абзац имеет выравнивание по левому краю вместо выравнивания по ширине.")
        paragraph2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Отсутствие отступа первой строки
        paragraph3 = doc.add_paragraph("У этого абзаца отсутствует отступ первой строки. " * 3)
        paragraph3.paragraph_format.first_line_indent = Cm(0)
        
        doc.save(file_path)
        
        self.document_files['multiple_errors'] = file_path
        print(f"Создан документ с множественными ошибками: {file_path}")
        return file_path
    
    def create_large_document(self):
        """
        Создает большой документ для тестирования производительности
        
        Returns:
            str: Путь к созданному файлу
        """
        file_path = TEST_DATA_DIR / "large_document.docx"
        
        doc = Document()
        # Добавляем 50 страниц текста
        for i in range(1, 51):
            doc.add_heading(f'Раздел {i}', level=1)
            for j in range(1, 11):
                doc.add_paragraph(f'Это параграф {j} в разделе {i}. ' * 20)
        
        doc.save(file_path)
        
        self.document_files['large'] = file_path
        print(f"Создан большой документ для тестирования производительности: {file_path}")
        return file_path
    
    def create_all_test_documents(self):
        """
        Создает все тестовые документы
        
        Returns:
            dict: Словарь с путями ко всем созданным файлам
        """
        self.create_empty_document()
        self.create_wrong_font_document()
        self.create_wrong_font_size_document()
        self.create_wrong_margins_document()
        self.create_wrong_line_spacing_document()
        self.create_multiple_errors_document()
        self.create_large_document()
        
        print(f"Все тестовые документы созданы в директории: {TEST_DATA_DIR}")
        return self.document_files


def parse_arguments():
    """
    Разбор аргументов командной строки
    
    Returns:
        argparse.Namespace: Аргументы командной строки
    """
    parser = argparse.ArgumentParser(description='Генерация тестовых DOCX-файлов')
    parser.add_argument('-a', '--all', action='store_true', help='Создать все тестовые документы')
    parser.add_argument('-e', '--empty', action='store_true', help='Создать пустой документ')
    parser.add_argument('-f', '--font', action='store_true', help='Создать документ с неправильным шрифтом')
    parser.add_argument('-s', '--size', action='store_true', help='Создать документ с неправильным размером шрифта')
    parser.add_argument('-m', '--margins', action='store_true', help='Создать документ с неправильными полями')
    parser.add_argument('-l', '--line-spacing', action='store_true', help='Создать документ с неправильным межстрочным интервалом')
    parser.add_argument('-u', '--multiple', action='store_true', help='Создать документ с множественными ошибками')
    parser.add_argument('-g', '--large', action='store_true', help='Создать большой документ для тестирования производительности')
    return parser.parse_args()


def main():
    """
    Основная функция
    """
    args = parse_arguments()
    generator = TestDocumentGenerator()
    
    if args.all:
        generator.create_all_test_documents()
        return 0
    
    # Создаем выбранные типы документов
    if args.empty:
        generator.create_empty_document()
    
    if args.font:
        generator.create_wrong_font_document()
    
    if args.size:
        generator.create_wrong_font_size_document()
    
    if args.margins:
        generator.create_wrong_margins_document()
    
    if args.line_spacing:
        generator.create_wrong_line_spacing_document()
    
    if args.multiple:
        generator.create_multiple_errors_document()
    
    if args.large:
        generator.create_large_document()
    
    # Если не указаны никакие параметры, создаем все документы по умолчанию
    if not any([args.all, args.empty, args.font, args.size, args.margins, args.line_spacing, args.multiple, args.large]):
        print("Не указаны параметры, создаем все тестовые документы по умолчанию...")
        generator.create_all_test_documents()
    
    return 0


if __name__ == "__main__":
    sys.exit(main()) 