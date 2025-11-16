import os
import tempfile
from docx import Document
from docx.shared import Pt
from app.services.document_corrector import DocumentCorrector

def test_pseudo_heading_promoted_and_spaced(tmp_path=None):
    # Arrange: создаем документ с псевдозаголовком обычным стилем
    tmp_dir = tmp_path if tmp_path else tempfile.mkdtemp()
    src = os.path.join(tmp_dir, 'source.docx')
    doc = Document()
    p1 = doc.add_paragraph("ГЛАВА 1. ВВЕДЕНИЕ")  # обычный текст, должен стать Heading 1
    p2 = doc.add_paragraph("Это абзац основного текста с неправильным межстрочным интервалом.")
    # Установим явно единичный интервал и отступы, чтобы проверить исправление
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_before = Pt(12)
    p2.paragraph_format.space_after = Pt(12)
    doc.save(src)

    # Act: прогоняем корректировку
    corrector = DocumentCorrector()
    out = corrector.correct_document(src, errors=None, out_path=os.path.join(tmp_dir, 'out.docx'))

    # Assert: открываем результат и проверяем
    fixed = Document(out)
    fp1 = fixed.paragraphs[0]
    fp2 = fixed.paragraphs[1]

    # Заголовок продвинут до Heading 1
    assert fp1.style.name.startswith('Heading'), "Первый параграф должен быть заголовком"

    # Межстрочный интервал в тексте — 1.5
    assert abs(fp2.paragraph_format.line_spacing - 1.5) < 1e-3, "Межстрочный интервал текста должен быть 1.5"

    # Интервалы до/после у обычного текста сброшены в ноль
    assert (fp2.paragraph_format.space_before is None or fp2.paragraph_format.space_before.pt == 0), "space_before должен быть 0"
    assert (fp2.paragraph_format.space_after is None or fp2.paragraph_format.space_after.pt == 0), "space_after должен быть 0"
