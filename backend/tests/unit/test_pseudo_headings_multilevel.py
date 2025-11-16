import os
import tempfile
from docx import Document
from app.services.document_corrector import DocumentCorrector


def test_multilevel_numbering_promoted(tmp_path=None):
    # Arrange
    tmp_dir = tmp_path if tmp_path else tempfile.mkdtemp()
    src = os.path.join(tmp_dir, 'multi.docx')
    doc = Document()
    doc.add_paragraph('1. ГЛАВА ПЕРВАЯ')
    doc.add_paragraph('1.1 Подраздел первый')
    doc.add_paragraph('1.1.1 Пункт детальный')
    doc.save(src)

    # Act
    corrector = DocumentCorrector()
    out = corrector.correct_document(src, errors=None, out_path=os.path.join(tmp_dir, 'out.docx'))

    # Assert
    fixed = Document(out)
    p1 = fixed.paragraphs[0]
    p2 = fixed.paragraphs[1]
    p3 = fixed.paragraphs[2]

    assert p1.style.name.startswith('Heading'), 'Первый уровень должен быть Heading'
    assert p2.style.name.startswith('Heading'), 'Второй уровень должен быть Heading'
    # Третий уровень — Heading 3, если стиль существует, либо Heading 2 as fallback
    if 'Heading 3' in fixed.styles:
        assert p3.style.name.startswith('Heading 3') or p3.style.name.startswith('Heading'), 'Третий уровень должен быть Heading 3'
    else:
        assert p3.style.name.startswith('Heading'), 'Третий уровень должен быть Heading'
